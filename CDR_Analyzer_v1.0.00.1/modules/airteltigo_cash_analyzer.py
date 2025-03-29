import os
import sys
import pandas as pd
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout, 
    QPushButton, QLabel, QFrame, QFileDialog, QMessageBox, 
    QTreeWidget, QTreeWidgetItem, QProgressBar, QScrollArea, 
    QLineEdit, QSplitter, QStackedWidget, QTabWidget
)
from PyQt5.QtCore import Qt, QThread, pyqtSignal
from PyQt5.QtGui import QPixmap
from docx import Document
from docx.shared import Inches
import datetime
import subprocess
import logging

# Suppress libpng warning
os.environ['QT_LOGGING_RULES'] = '*.warning=false'

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('airteltigo_cash_analyzer.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# Color Constants
BACKGROUND_COLOR = "#1E1E2F"  # Dark background
SIDEBAR_COLOR = "#2C2C3A"     # Sidebar background
CONTENT_BG_COLOR = "#1e3b75"  # Content area background
TEXT_COLOR = "#FFFFFF"        # White text
BUTTON_COLOR = "#213d79"      # Button color
BUTTON_TEXT_COLOR = "#FFFFFF" # Button text color
HIGHLIGHT_COLOR = "#CC0000"   # Hover color
ERROR_COLOR = "#FF6B6B"       # Error color
SUCCESS_COLOR = "#4CAF50"     # Success color
BORDER_COLOR = "#4a3a4a"      # Border color
ACCENT_COLOR = "#1e3b75"      # Accent color
ACCENT_TEXT = "#FFFFFF"       # Accent text color

APP_STYLESHEET = f"""
    /* Main window */
    QMainWindow {{
        background-color: {BACKGROUND_COLOR};
        color: {TEXT_COLOR};
    }}
    
    /* Sidebar */
    QFrame#sidebar {{
        background-color: {SIDEBAR_COLOR};
        border-right: 1px solid {BORDER_COLOR};
    }}
    
    /* Content area */
    QWidget#content_area {{
        background-color: {BACKGROUND_COLOR};
        padding: 20px;
    }}
    
    /* Tree widget */
    QTreeWidget {{
        background-color: {CONTENT_BG_COLOR};
        color: {TEXT_COLOR};
        border: 1px solid {BORDER_COLOR};
        font-size: 12px;
    }}
    
    QTreeWidget::item {{
        padding: 5px;
    }}
    
    QHeaderView::section {{
        background-color: {CONTENT_BG_COLOR};
        color: {TEXT_COLOR};
        padding: 5px;
        border: none;
    }}
    
    /* Standard Buttons */
    QPushButton {{
        background-color: {BUTTON_COLOR};
        color: {BUTTON_TEXT_COLOR};
        font-size: 14px;
        font-weight: bold;
        padding: 8px 12px;
        border: none;
        border-radius: 5px;
        min-width: 100px;
        min-height: 40px;
    }}
    
    QPushButton:hover {{
        background-color: {HIGHLIGHT_COLOR};
    }}
    
    /* Sidebar Buttons */
    QFrame#sidebar QPushButton {{
        min-width: 180px;
        min-height: 45px;
    }}
    
    /* Important Buttons */
    QPushButton.important {{
        min-width: 150px;
        padding: 10px 15px;
    }}
    
    /* Reset Buttons */
    QPushButton.reset {{
        background-color: {ERROR_COLOR};
        color: white;
    }}
    
    QPushButton.reset:hover {{
        background-color: #C0392B;
    }}
    
    /* Progress bar */
    QProgressBar {{
        background: {CONTENT_BG_COLOR};
        border: 1px solid {BORDER_COLOR};
        text-align: center;
        color: {TEXT_COLOR};
    }}
    
    QProgressBar::chunk {{
        background-color: {ACCENT_COLOR};
    }}
    
    /* Line edits */
    QLineEdit {{
        background: {CONTENT_BG_COLOR};
        border: 1px solid {BORDER_COLOR};
        padding: 5px;
        color: {TEXT_COLOR};
    }}
    
    /* Labels */
    QLabel {{
        color: {TEXT_COLOR};
    }}
    
    QLabel#file_label {{
        qproperty-alignment: AlignCenter;
        font-size: 14px;
    }}
    
    QLabel#file_label[error=true] {{
        color: {ERROR_COLOR};
    }}
    
    QLabel#file_label[success=true] {{
        color: {SUCCESS_COLOR};
    }}
    
    /* Scroll areas */
    QScrollArea {{
        border: none;
    }}
    
    QScrollBar:vertical {{
        background: {CONTENT_BG_COLOR};
        width: 10px;
    }}
    
    QScrollBar::handle:vertical {{
        background: {BORDER_COLOR};
        min-height: 20px;
    }}
    
    /* Logo frame - matches sidebar background */
    QFrame#logo_frame {{
        background-color: {SIDEBAR_COLOR};
        padding: 5px;
        margin-top: 10px;
        border: none;
    }}
    
    /* Tab widget */
    QTabWidget::pane {{
        border: 1px solid {BORDER_COLOR};
    }}
    
    QTabBar::tab {{
        background: {CONTENT_BG_COLOR};
        color: {TEXT_COLOR};
        padding: 8px;
        border: 1px solid {BORDER_COLOR};
        border-bottom: none;
    }}
    
    QTabBar::tab:selected {{
        background: {ACCENT_COLOR};
        color: {ACCENT_TEXT};
    }}
    
    /* Sidebar toggle button */
    QPushButton#toggleSidebar {{
        background-color: transparent;
        color: white;
        font-size: 16px;
        font-weight: bold;
        border: none;
        padding: 0;
    }}
    
    QPushButton#toggleSidebar:hover {{
        background-color: rgba(255, 255, 255, 0.1);
        border-radius: 3px;
    }}
"""

class AnalysisThread(QThread):
    update_progress = pyqtSignal(str, int)
    analysis_complete = pyqtSignal(list, dict)
    error_occurred = pyqtSignal(str)

    def __init__(self, df):
        super().__init__()
        self.df = df
        self.visualizations = {}
        self.insights = []

    def run(self):
        try:
            if self.df is None:
                self.error_occurred.emit("No data loaded for analysis")
                return

            results = []
            total_steps = 10
            current_step = 0

            # Basic transaction analysis
            self.update_progress.emit("Analyzing transaction patterns...", int(current_step/total_steps*100))
            
            # 1. Total Transactions
            total_transactions = len(self.df)
            results.append(("Total Transactions", f"{total_transactions}"))
            
            # 2. Total Deposits (Paid In)
            total_income = self.df['Paid In'].sum()
            results.append(("Total Deposits", f"GH₵ {total_income:.2f}"))
            
            # 3. Total Withdrawals
            total_expenses = self.df['Withdrawn'].sum()
            results.append(("Total Withdrawals", f"GH₵ {total_expenses:.2f}"))
            
            # 4. Net Cash Flow
            net_balance = total_income - total_expenses
            results.append(("Net Cash Flow", f"GH₵ {net_balance:.2f}"))
            
            # 5. Average Deposit Amount
            avg_deposit = self.df[self.df['Paid In'] > 0]['Paid In'].mean()
            results.append(("Average Deposit Amount", f"GH₵ {avg_deposit:.2f}" if not pd.isna(avg_deposit) else "N/A"))
            
            # 6. Average Withdrawal Amount
            avg_withdrawal = self.df[self.df['Withdrawn'] > 0]['Withdrawn'].mean()
            results.append(("Average Withdrawal Amount", f"GH₵ {avg_withdrawal:.2f}" if not pd.isna(avg_withdrawal) else "N/A"))
            
            current_step += 1
            self.update_progress.emit("Analyzing transaction status...", int(current_step/total_steps*100))

            # 7. Transaction Status
            status_counts = self.df['Transaction Status'].value_counts()
            results.append(("Transaction Status", "Count"))
            
            for status, count in status_counts.items():
                results.append((f"{status}", f"{count}"))
                
            # Generate transaction status pie chart
            self.generate_pie_chart(
                status_counts,
                "Transaction Status Distribution",
                "transaction_status.png"
            )
            self.visualizations['transaction_status'] = "transaction_status.png"
            self.insights.append(f"Transaction status: {', '.join([f'{k} ({v})' for k,v in status_counts.items()])}")

            current_step += 1
            self.update_progress.emit("Analyzing transaction parties...", int(current_step/total_steps*100))

            # 8. Top Counterparties
            top_parties = self.df['Opposite Party'].value_counts().head(5)
            results.append(("Top 5 Counterparties", "Count"))
            
            for party, count in top_parties.items():
                results.append((f"{party}", f"{count}"))
                
            # Generate top parties bar chart
            self.generate_bar_chart(
                top_parties,
                "Top 5 Counterparties",
                "Counterparty",
                "Number of Transactions",
                "top_parties.png"
            )
            self.visualizations['top_parties'] = "top_parties.png"
            self.insights.append(f"Most frequent counterparty: {top_parties.index[0]} with {top_parties.iloc[0]} transactions")

            current_step += 1
            self.update_progress.emit("Analyzing transaction amounts...", int(current_step/total_steps*100))

            # 9. Transaction Amount Distribution
            transaction_amounts = pd.concat([
                self.df['Paid In'], 
                self.df['Withdrawn']
            ])
            self.generate_histogram(
                transaction_amounts,
                "Transaction Amount Distribution",
                "Amount (GH₵)",
                "Number of Transactions",
                "amount_distribution.png",
                bins=20
            )
            self.visualizations['amount_dist'] = "amount_distribution.png"
            self.insights.append(f"Average transaction amount: GH₵ {transaction_amounts.mean():.2f}")

            current_step += 1
            self.update_progress.emit("Generating time-based analysis...", int(current_step/total_steps*100))

            # 10. Time-based analysis
            if 'Completion Time' in self.df.columns:
                try:
                    self.df['DATE'] = pd.to_datetime(self.df['Completion Time']).dt.date
                    self.df['HOUR'] = pd.to_datetime(self.df['Completion Time']).dt.hour
                    self.df['DAY_OF_WEEK'] = pd.to_datetime(self.df['Completion Time']).dt.day_name()
                    
                    # Hourly distribution
                    hourly_counts = self.df['HOUR'].value_counts().sort_index()
                    self.generate_bar_chart(
                        hourly_counts,
                        "Hourly Transaction Distribution",
                        "Hour of Day",
                        "Number of Transactions",
                        "hourly_distribution.png"
                    )
                    self.visualizations['hourly'] = "hourly_distribution.png"
                    
                    peak_hour = hourly_counts.idxmax()
                    self.insights.append(f"Peak transaction hour: {peak_hour}:00 with {hourly_counts.max()} transactions")
                    
                    # Day of week distribution
                    day_counts = self.df['DAY_OF_WEEK'].value_counts()
                    self.generate_bar_chart(
                        day_counts,
                        "Day of Week Distribution",
                        "Day",
                        "Number of Transactions",
                        "day_distribution.png"
                    )
                    self.visualizations['daily'] = "day_distribution.png"
                    
                    peak_day = day_counts.idxmax()
                    self.insights.append(f"Most active day: {peak_day} with {day_counts.max()} transactions")
                except Exception as e:
                    logger.error(f"Time analysis error: {str(e)}")

            current_step += 1
            self.update_progress.emit("Finalizing results...", int(current_step/total_steps*100))

            self.analysis_complete.emit(results, {
                'visualizations': self.visualizations,
                'insights': self.insights
            })

        except Exception as e:
            logger.error(f"Analysis error: {str(e)}")
            self.error_occurred.emit(f"Analysis failed: {str(e)}")

    def generate_pie_chart(self, data, title, filename):
        """Generate a pie chart visualization"""
        plt.figure(figsize=(8, 8))
        data.plot(kind='pie', autopct='%1.1f%%', startangle=90)
        plt.title(title)
        plt.ylabel('')
        plt.savefig(filename, dpi=100, bbox_inches='tight')
        plt.close()

    def generate_histogram(self, data, title, xlabel, ylabel, filename, bins=10):
        """Generate a histogram visualization"""
        plt.figure(figsize=(10, 6))
        plt.hist(data, bins=bins, edgecolor='black')
        plt.title(title)
        plt.xlabel(xlabel)
        plt.ylabel(ylabel)
        plt.grid(True)
        plt.savefig(filename, dpi=100, bbox_inches='tight')
        plt.close()

    def generate_bar_chart(self, data, title, xlabel, ylabel, filename):
        """Generate a bar chart visualization"""
        plt.figure(figsize=(10, 6))
        data.plot(kind='bar', color=BUTTON_COLOR)
        plt.title(title)
        plt.xlabel(xlabel)
        plt.ylabel(ylabel)
        plt.grid(True)
        plt.savefig(filename, dpi=100, bbox_inches='tight')
        plt.close()


class AirtelTigoCashAnalyzer(QMainWindow):
    def __init__(self, back_to_home_callback):
        super().__init__()
        self.back_to_home_callback = back_to_home_callback
        self.setWindowTitle("AirtelTigo Cash Transaction Analyzer")
        self.setGeometry(100, 100, 1200, 800)
        self.setStyleSheet(APP_STYLESHEET)
        
        # Initialize variables
        self.filename = None
        self.df = None
        self.doc_path = None
        self.analysis_thread = None
        self.visualizations = {}
        self.insights = []
        
        # Create main widgets
        self.init_ui()
        self.create_home_page()
    
    def init_ui(self):
        """Initialize main UI components"""
        self.central_widget = QWidget()
        self.setCentralWidget(self.central_widget)
        self.main_layout = QHBoxLayout(self.central_widget)
        self.main_layout.setContentsMargins(0, 0, 0, 0)
        self.main_layout.setSpacing(0)

        # Sidebar
        self.sidebar = QFrame()
        self.sidebar.setObjectName("sidebar")
        self.sidebar_width = 240  # Store the original width
        self.sidebar.setFixedWidth(self.sidebar_width)
        self.sidebar_layout = QVBoxLayout(self.sidebar)
        self.sidebar_layout.setAlignment(Qt.AlignTop)
        self.sidebar_layout.setSpacing(10)
        self.sidebar_layout.setContentsMargins(10, 20, 10, 10)
        
        # Add toggle button at the top of sidebar
        self.toggle_sidebar_btn = QPushButton("◄")
        self.toggle_sidebar_btn.setObjectName("toggleSidebar")
        self.toggle_sidebar_btn.setFixedSize(30, 30)
        self.toggle_sidebar_btn.clicked.connect(self.toggle_sidebar)
        self.sidebar_layout.addWidget(self.toggle_sidebar_btn, 0, Qt.AlignLeft | Qt.AlignTop)
        
        # Add buttons to sidebar
        self.create_sidebar_buttons()
        self.sidebar_layout.addStretch()
        self.add_sidebar_logo()
        
        self.main_layout.addWidget(self.sidebar)
        
        # Create content area
        self.content_stack = QStackedWidget()
        self.main_layout.addWidget(self.content_stack)
    
    def toggle_sidebar(self):
        """Toggle sidebar visibility"""
        if self.sidebar.width() > 50:  # If sidebar is expanded
            # Minimize sidebar
            self.sidebar.setFixedWidth(50)
            self.toggle_sidebar_btn.setText("►")
            
            # Hide all widgets except the toggle button
            for i in range(1, self.sidebar_layout.count()):
                item = self.sidebar_layout.itemAt(i)
                if item.widget():
                    item.widget().hide()
        else:
            # Maximize sidebar
            self.sidebar.setFixedWidth(self.sidebar_width)
            self.toggle_sidebar_btn.setText("◄")
            
            # Show all widgets
            for i in range(1, self.sidebar_layout.count()):
                item = self.sidebar_layout.itemAt(i)
                if item.widget():
                    item.widget().show()
    
    def create_sidebar_buttons(self):
        """Create sidebar navigation buttons"""
        buttons = [
            ("Load File", self.load_file, "standard"),
            ("Analyze Data", self.create_analysis_page, "standard"),
            ("Export Report", self.export_report, "standard"),
            ("Back to Home", self.back_to_home, "warning")
        ]
        
        for text, callback, btn_type in buttons:
            btn = QPushButton(text)
            if btn_type == "warning":
                btn.setProperty("class", "reset")
            btn.clicked.connect(callback)
            self.sidebar_layout.addWidget(btn)
            
            if text == "Export Report":
                self.export_button = btn
                self.export_button.setEnabled(False)
    
    def add_sidebar_logo(self):
        """Add AirtelTigo logo to sidebar with matching background"""
        logo_frame = QFrame()
        logo_frame.setObjectName("logo_frame")
        logo_frame.setStyleSheet(f"background-color: {SIDEBAR_COLOR}; border: none;")
        logo_layout = QVBoxLayout(logo_frame)
        logo_layout.setAlignment(Qt.AlignCenter)
        logo_layout.setContentsMargins(5, 5, 5, 5)
        
        logo_label = QLabel()
        logo_path = os.path.join("assets", "airteltigo_cash.png")
        if os.path.exists(logo_path):
            logo_pixmap = QPixmap(logo_path)
            if not logo_pixmap.isNull():
                # Scale logo to appropriate size
                logo_pixmap = logo_pixmap.scaled(170, 170, Qt.KeepAspectRatio, Qt.SmoothTransformation)
                logo_label.setPixmap(logo_pixmap)
                logo_label.setAlignment(Qt.AlignCenter)
        else:
            # Fallback if logo not found
            logo_label.setText("AirtelTigo Cash")
            logo_label.setStyleSheet("""
                font-size: 16px; 
                font-weight: bold; 
                color: white;
            """)
        
        logo_layout.addWidget(logo_label)
        self.sidebar_layout.addWidget(logo_frame)
    
    def back_to_home(self):
        """Reset the app and return to home screen."""
        # Reset all variables
        self.filename = None
        self.df = None
        self.doc_path = None
        self.visualizations = {}
        self.insights = []
        
        # Clear the content
        self.clear_content()
        
        # Reset file label
        self.file_label.setText("📂 No file selected")
        self.file_label.setProperty("error", True)
        self.file_label.style().polish(self.file_label)
        
        # Disable export button
        self.export_button.setEnabled(False)
        
        # Return to home page
        self.content_stack.setCurrentWidget(self.home_page)
        
        # If we have a callback, use it to exit completely
        if self.back_to_home_callback:
            self.back_to_home_callback()
    
    def clear_content(self):
        """Clear the content area"""
        if hasattr(self, 'analysis_page'):
            layout = self.analysis_page.layout()
            if layout:
                while layout.count():
                    item = layout.takeAt(0)
                    widget = item.widget()
                    if widget:
                        widget.deleteLater()
    
    def create_home_page(self):
        """Create the home page"""
        if hasattr(self, 'home_page'):
            self.content_stack.setCurrentWidget(self.home_page)
            return
        
        self.home_page = QWidget()
        home_layout = QVBoxLayout(self.home_page)
        home_layout.setAlignment(Qt.AlignCenter)
        
        title = QLabel("AirtelTigo Cash Transaction Analyzer")
        title.setStyleSheet("font-size: 20px; font-weight: bold;")
        title.setAlignment(Qt.AlignCenter)
        home_layout.addWidget(title)
        
        self.file_label = QLabel("📂 No file selected")
        self.file_label.setObjectName("file_label")
        self.file_label.setProperty("error", True)
        self.file_label.style().polish(self.file_label)
        home_layout.addWidget(self.file_label)
        
        home_layout.addStretch()
        self.content_stack.addWidget(self.home_page)
        self.content_stack.setCurrentWidget(self.home_page)
    
    def create_analysis_page(self):
        """Create the analysis page with results"""
        if self.df is None:
            QMessageBox.warning(self, "Error", "Please load a file first!")
            return
        
        if hasattr(self, 'analysis_page'):
            self.content_stack.setCurrentWidget(self.analysis_page)
            return
        
        self.analysis_page = QWidget()
        analysis_layout = QVBoxLayout(self.analysis_page)
        
        title = QLabel("AirtelTigo Cash Analysis Report")
        title.setStyleSheet("font-size: 18px; font-weight: bold;")
        title.setAlignment(Qt.AlignCenter)
        analysis_layout.addWidget(title)
        
        search_frame = QFrame()
        search_layout = QHBoxLayout(search_frame)
        
        self.search_input = QLineEdit()
        self.search_input.setPlaceholderText("Search...")
        self.search_input.textChanged.connect(self.filter_treeview)
        search_layout.addWidget(self.search_input)
        
        reset_btn = QPushButton("⭮ Reset")
        reset_btn.setProperty("class", "reset")
        reset_btn.clicked.connect(self.reset_filters)
        search_layout.addWidget(reset_btn)
        
        analysis_layout.addWidget(search_frame)
        
        # Create tab widget for results and visualizations
        self.tab_widget = QTabWidget()
        
        # Results tab
        results_tab = QWidget()
        results_layout = QVBoxLayout(results_tab)
        
        self.tree = QTreeWidget()
        self.tree.setHeaderLabels(["Metric", "Value"])
        self.tree.setColumnWidth(0, 300)
        self.tree.setColumnWidth(1, 500)
        results_layout.addWidget(self.tree)
        
        # Insights section
        insights_label = QLabel("Key Insights")
        insights_label.setStyleSheet("font-size: 16px; font-weight: bold;")
        results_layout.addWidget(insights_label)
        
        self.insights_text = QLabel("Analysis insights will appear here...")
        self.insights_text.setWordWrap(True)
        results_layout.addWidget(self.insights_text)
        
        self.tab_widget.addTab(results_tab, "Results")
        
        # Visualizations tab
        viz_tab = QWidget()
        viz_layout = QVBoxLayout(viz_tab)
        
        self.viz_scroll = QScrollArea()
        self.viz_scroll.setWidgetResizable(True)
        self.viz_container = QWidget()
        self.viz_layout = QVBoxLayout(self.viz_container)
        
        # Placeholder for visualizations
        self.viz_placeholder = QLabel("Visualizations will appear here after analysis")
        self.viz_placeholder.setAlignment(Qt.AlignCenter)
        self.viz_layout.addWidget(self.viz_placeholder)
        
        self.viz_scroll.setWidget(self.viz_container)
        viz_layout.addWidget(self.viz_scroll)
        
        self.tab_widget.addTab(viz_tab, "Visualizations")
        
        analysis_layout.addWidget(self.tab_widget)
        
        self.progress_bar = QProgressBar()
        self.progress_bar.setRange(0, 100)
        analysis_layout.addWidget(self.progress_bar)
        
        self.status_label = QLabel("Analyzing data... Please wait")
        self.status_label.setAlignment(Qt.AlignCenter)
        analysis_layout.addWidget(self.status_label)
        
        self.content_stack.addWidget(self.analysis_page)
        self.content_stack.setCurrentWidget(self.analysis_page)
        
        self.start_analysis()
    
    def start_analysis(self):
        """Start the analysis in a background thread"""
        if self.analysis_thread and self.analysis_thread.isRunning():
            self.analysis_thread.terminate()
        
        self.analysis_thread = AnalysisThread(self.df)
        self.analysis_thread.update_progress.connect(self.update_progress_status)
        self.analysis_thread.analysis_complete.connect(self.display_results)
        self.analysis_thread.error_occurred.connect(self.show_error)
        self.analysis_thread.start()
    
    def update_progress_status(self, message, percent):
        """Update progress status"""
        self.status_label.setText(message)
        self.progress_bar.setValue(percent)
    
    def display_results(self, results, analysis_data):
        """Display analysis results and visualizations"""
        self.tree.clear()
        self.visualizations = analysis_data['visualizations']
        self.insights = analysis_data['insights']
        
        # Display results in tree widget
        for category, details in results:
            item = QTreeWidgetItem([str(category), str(details)])
            self.tree.addTopLevelItem(item)
        
        # Display insights
        self.insights_text.setText("\n• ".join([""] + self.insights))
        
        # Display visualizations
        self.show_visualizations()
        
        # Enable export button
        self.export_button.setEnabled(True)
        
        self.status_label.setText("Analysis Complete")
    
    def show_visualizations(self):
        """Display all generated visualizations"""
        # Clear previous visualizations
        for i in reversed(range(self.viz_layout.count())): 
            self.viz_layout.itemAt(i).widget().setParent(None)
        
        # Remove placeholder
        if hasattr(self, 'viz_placeholder'):
            self.viz_layout.removeWidget(self.viz_placeholder)
            self.viz_placeholder.deleteLater()
            del self.viz_placeholder
        
        # Add each visualization
        for viz_type, viz_path in self.visualizations.items():
            if os.path.exists(viz_path):
                # Add title
                title = QLabel(self.get_viz_title(viz_type))
                title.setStyleSheet("font-size: 14px; font-weight: bold;")
                self.viz_layout.addWidget(title)
                
                # Add image
                pixmap = QPixmap(viz_path)
                if not pixmap.isNull():
                    label = QLabel()
                    label.setPixmap(pixmap.scaled(800, 600, Qt.KeepAspectRatio, Qt.SmoothTransformation))
                    label.setAlignment(Qt.AlignCenter)
                    self.viz_layout.addWidget(label)
        
        self.viz_layout.addStretch()
    
    def get_viz_title(self, viz_type):
        """Get display title for visualization type"""
        titles = {
            'transaction_status': "Transaction Status Distribution",
            'top_parties': "Top 5 Counterparties",
            'amount_dist': "Transaction Amount Distribution",
            'hourly': "Hourly Transaction Distribution",
            'daily': "Daily Transaction Distribution"
        }
        return titles.get(viz_type, viz_type.replace('_', ' ').title())
    
    def filter_treeview(self):
        """Filter tree items based on search text"""
        search_text = self.search_input.text().lower()
        
        for i in range(self.tree.topLevelItemCount()):
            item = self.tree.topLevelItem(i)
            category = item.text(0).lower()
            details = item.text(1).lower()
            item.setHidden(search_text not in category and search_text not in details)
    
    def reset_filters(self):
        """Reset all filters"""
        self.search_input.clear()
        for i in range(self.tree.topLevelItemCount()):
            self.tree.topLevelItem(i).setHidden(False)
    
    def load_file(self):
        """Load Excel file with AirtelTigo Cash format"""
        filename, _ = QFileDialog.getOpenFileName(
            self, "Open AirtelTigo Cash File", "", "Excel Files (*.xlsx *.xls);;CSV Files (*.csv)"
        )
        
        if filename:
            try:
                if filename.endswith('.csv'):
                    self.df = pd.read_csv(filename)
                else:
                    self.df = pd.read_excel(filename, engine='openpyxl')

                # Normalize column names (remove extra spaces, etc.)
                self.df.columns = [col.strip() for col in self.df.columns]

                # Check for required columns in AirtelTigo format
                required_columns = ["Paid In", "Withdrawn", "Balance", "Opposite Party"]
                missing_columns = [col for col in required_columns if col not in self.df.columns]
                
                if missing_columns:
                    QMessageBox.critical(
                        self, 
                        "Error", 
                        f"File doesn't match AirtelTigo Cash format.\nMissing columns: {', '.join(missing_columns)}"
                    )
                    self.df = None
                    return

                # Clean data - replace empty strings with 0 for numeric columns
                numeric_cols = ['Paid In', 'Withdrawn', 'Balance']
                for col in numeric_cols:
                    if col in self.df.columns:
                        self.df[col] = pd.to_numeric(self.df[col], errors='coerce').fillna(0)

                self.filename = filename
                self.file_label.setText(f"📂 Loaded: {os.path.basename(filename)}")
                self.file_label.setProperty("error", False)
                self.file_label.setProperty("success", True)
                self.file_label.style().polish(self.file_label)
                
                # Disable export button until analysis is complete
                self.export_button.setEnabled(False)
                
                QMessageBox.information(self, "Success", "AirtelTigo Cash file loaded successfully!")
                
            except Exception as e:
                logger.error(f"File loading error: {str(e)}")
                QMessageBox.critical(self, "Error", f"Failed to load file:\n{str(e)}")
                self.df = None
                self.file_label.setText("📂 Error loading file")
                self.file_label.setProperty("error", True)
                self.file_label.setProperty("success", False)
                self.file_label.style().polish(self.file_label)
    
    def export_report(self):
        """Export report to Word document with AirtelTigo branding"""
        if not hasattr(self, 'tree') or self.tree.topLevelItemCount() == 0:
            QMessageBox.warning(self, "Error", "No analysis results to export!")
            return
        
        filename, _ = QFileDialog.getSaveFileName(
            self, "Save AirtelTigo Cash Report", "", "Word Documents (*.docx)"
        )
        
        if not filename:
            return
        
        try:
            doc = Document()
            
            # Add title and metadata with AirtelTigo branding
            doc.add_heading('AirtelTigo Cash Transaction Analysis Report', 0)
            current_time = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            doc.add_paragraph(f"Report generated on: {current_time}")
            doc.add_paragraph(f"Source file: {os.path.basename(self.filename)}")
            doc.add_paragraph("\n")

            # Add summary statistics
            doc.add_heading('Summary Statistics', level=1)
            for i in range(self.tree.topLevelItemCount()):
                item = self.tree.topLevelItem(i)
                doc.add_paragraph(f"{item.text(0)}: {item.text(1)}")
            doc.add_paragraph("\n")

            # Add insights
            doc.add_heading('Key Insights', level=1)
            for insight in self.insights:
                doc.add_paragraph(insight, style='List Bullet')
            doc.add_paragraph("\n")

            # Add visualizations
            doc.add_heading('Visualizations', level=1)
            
            # Add each visualization that exists in both UI and filesystem
            viz_order = ['transaction_status', 'top_parties', 'amount_dist', 'hourly', 'daily']
            for viz_type in viz_order:
                if viz_type in self.visualizations and os.path.exists(self.visualizations[viz_type]):
                    doc.add_heading(self.get_viz_title(viz_type), level=2)
                    doc.add_picture(self.visualizations[viz_type], width=Inches(6))
                    doc.add_paragraph("\n")

            # Save document
            doc.save(filename)

            try:
                os.startfile(filename)  # For Windows
            except:
                subprocess.run(["open", filename])  # For macOS
                subprocess.run(["xdg-open", filename])  # For Linux

            QMessageBox.information(self, "Success", f"AirtelTigo Cash report exported to {filename}")

        except Exception as e:
            logger.error(f"Report export error: {str(e)}")
            QMessageBox.critical(self, "Error", f"Failed to export report:\n{str(e)}")
    
    def show_error(self, message):
        """Show error message"""
        QMessageBox.critical(self, "Error", message)
        self.status_label.setText("Analysis Failed")
        self.progress_bar.setValue(0)


if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = AirtelTigoCashAnalyzer(lambda: None)  # Placeholder callback for standalone testing
    window.show()
    sys.exit(app.exec_())