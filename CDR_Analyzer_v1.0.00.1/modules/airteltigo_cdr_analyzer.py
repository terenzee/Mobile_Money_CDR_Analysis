import os
import sys
import pandas as pd
import folium
import webbrowser
import matplotlib.pyplot as plt
import matplotlib
matplotlib.use('Agg')
from folium.plugins import HeatMap
from geopy.geocoders import Nominatim
from geopy.exc import GeocoderTimedOut, GeocoderUnavailable
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout, QPushButton,
    QLabel, QFrame, QFileDialog, QMessageBox, QTreeWidget, QTreeWidgetItem,
    QProgressBar, QScrollArea, QLineEdit, QSplitter, QSizePolicy, QStackedWidget,
    QTabWidget
)
from PyQt5.QtCore import Qt, QThread, pyqtSignal, QUrl, QEasingCurve, QPropertyAnimation, QPoint
from PyQt5.QtGui import QPixmap, QIcon
import datetime
import math
import logging
import numpy as np
from collections import Counter

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('airteltigo_cdr_analyzer.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# Color Constants
BACKGROUND_COLOR = "#1E1E2F"
SIDEBAR_COLOR = "#2C2C3A"
CONTENT_BG_COLOR = "#1e3b75"
TEXT_COLOR = "#FFFFFF"
BUTTON_COLOR = "#213d79"
BUTTON_TEXT_COLOR = "#FFFFFF"
HIGHLIGHT_COLOR = "#CC0000"
ERROR_COLOR = "#FF6B6B"
SUCCESS_COLOR = "#4CAF50"
BORDER_COLOR = "#4a3a4a"
ACCENT_COLOR = "#1e3b75"
ACCENT_TEXT = "#FFFFFF"

class AnalysisThread(QThread):
    update_progress = pyqtSignal(str, int)
    analysis_complete = pyqtSignal(list, dict)
    error_occurred = pyqtSignal(str)

    def __init__(self, df):
        super().__init__()
        self.df = df
        self.geolocator = Nominatim(user_agent="airteltigo_cdr_analyzer")
        self.geolocation_cache = {}
        self.location_counts = None
        self.imei_counts = None
        self.visualizations = {}
        self.insights = []

    def run(self):
        try:
            if self.df is None:
                self.error_occurred.emit("No data loaded for analysis")
                return

            results = []
            total_steps = 15
            current_step = 0

            # Basic call analysis
            self.update_progress.emit("Analyzing call patterns...", int(current_step/total_steps*100))
            
            self.df['Outgoing'] = self.df['Outgoing'].fillna(0).astype('Int64')
            self.df['Incoming'] = self.df['Incoming'].fillna(0).astype('Int64')
            
            calls = self.df
            total_calls = len(calls)
            unique_numbers = pd.concat([
                calls["Owner Number"], 
                calls["Outgoing"].replace(0, pd.NA).dropna(), 
                calls["Incoming"].replace(0, pd.NA).dropna()
            ]).nunique()
            
            outgoing_contacts = calls["Outgoing"].value_counts().head(5)
            incoming_contacts = calls["Incoming"].value_counts().head(5)
            
            call_durations = calls[calls["Call Type"] != "SMS"]["Duration"]
            total_duration = call_durations.sum()
            avg_duration = call_durations.mean()

            results.extend([
                ("Total Records", f"{total_calls}"),
                ("Unique Phone Numbers", f"{unique_numbers}"),
                ("Total Call Duration (sec)", f"{total_duration}"),
                ("Average Call Duration (sec)", f"{avg_duration:.2f}"),
                ("Top Outgoing Contacts", "Frequency")
            ])

            for contact, freq in outgoing_contacts.items():
                results.append((f"{contact}", f"{freq}"))
            
            results.append(("Top Incoming Contacts", "Frequency"))
            for contact, freq in incoming_contacts.items():
                results.append((f"{contact}", f"{freq}"))
            
            current_step += 1
            self.update_progress.emit("Analyzing device information...", int(current_step/total_steps*100))

            # Device analysis
            self.imei_counts = calls["IMEI"].value_counts().sort_values(ascending=False)
            unique_imeis = len(self.imei_counts)
            unique_imsis = calls["IMSI"].nunique()
            cell_details = calls["Cell Details"].value_counts().head(5)
            
            results.extend([
                ("Unique IMEIs", f"{unique_imeis}"),
                ("IMEI Frequency Analysis", "Records")
            ])
            
            for i, (imei, count) in enumerate(self.imei_counts.head(5).items(), 1):
                results.append((f"{imei}", f"{count}"))
            
            results.extend([
                ("Unique IMSIs", f"{unique_imsis}"),
                ("Most Used Cell Towers", "Records")
            ])
            
            for tower, count in cell_details.items():
                results.append((f"{tower}", f"{count}"))

            current_step += 1
            self.update_progress.emit("Analyzing geolocations...", int(current_step/total_steps*100))

            # Geolocation analysis
            self.location_counts = calls.groupby(["Latitude", "Longitude"]).size()\
                                 .reset_index(name="count")\
                                 .sort_values("count", ascending=False)
            
            results.append(("Geolocation Frequency Analysis", "Records"))
            for i, (_, row) in enumerate(self.location_counts.head(5).iterrows(), 1):
                lat = f"{row['Latitude']:.6f}"
                lon = f"{row['Longitude']:.6f}"
                results.append((f"Lat: {lat}, Lon: {lon}", f"{row['count']}"))
            
            current_step += 1
            self.update_progress.emit("Analyzing temporal patterns...", int(current_step/total_steps*100))

            # Temporal Analysis
            calls['Event Date & Time'] = pd.to_datetime(calls['Event Date & Time'])
            calls['Hour'] = calls['Event Date & Time'].dt.hour
            calls['DayOfWeek'] = calls['Event Date & Time'].dt.day_name()
            calls['Period'] = calls['Hour'].apply(self.get_time_period)
            
            period_counts = calls['Period'].value_counts()
            most_common_period = period_counts.idxmax()
            
            day_counts = calls['DayOfWeek'].value_counts()
            most_common_day = day_counts.idxmax()

            results.extend([
                ("Most Active Period", most_common_period),
                ("Most Active Day", most_common_day)
            ])

            current_step += 1
            self.update_progress.emit("Generating call type distribution...", int(current_step/total_steps*100))

            # Call type distribution
            call_type_dist = calls['Call Type'].value_counts()
            self.generate_pie_chart(
                call_type_dist, 
                "Call Type Distribution", 
                "call_type_distribution.png"
            )
            self.visualizations['call_type'] = "call_type_distribution.png"
            self.insights.append(f"Call types: {', '.join([f'{k} ({v} calls)' for k,v in call_type_dist.items()])}")

            current_step += 1
            self.update_progress.emit("Generating duration histogram...", int(current_step/total_steps*100))

            # Call duration histogram
            call_durations = calls[calls['Call Type'] != 'SMS']['Duration']
            self.generate_histogram(
                call_durations,
                "Call Duration Distribution",
                "Duration (seconds)",
                "Number of Calls",
                "duration_distribution.png",
                bins=20
            )
            self.visualizations['duration'] = "duration_distribution.png"
            self.insights.append(f"Average call duration: {avg_duration:.2f} seconds")

            current_step += 1
            self.update_progress.emit("Generating geolocation plot...", int(current_step/total_steps*100))

            # Generate geolocation visualization
            self.generate_geolocation_plot(calls)
            self.visualizations['geolocation'] = "geolocation_plot.png"
            
            if not self.location_counts.empty:
                top_location = self.location_counts.iloc[0]
                self.insights.append(f"Most frequent location: Lat {top_location['Latitude']:.4f}, Lon {top_location['Longitude']:.4f} with {top_location['count']} calls")

            current_step += 1
            self.update_progress.emit("Generating time distribution charts...", int(current_step/total_steps*100))

            # Generate time distribution charts
            self.generate_time_distribution_charts(calls)
            self.visualizations['hourly'] = "hourly_distribution.png"
            self.visualizations['daily'] = "day_distribution.png"
            
            hour_dist = calls['Hour'].value_counts()
            peak_hour = hour_dist.idxmax()
            self.insights.append(f"Peak calling hour: {peak_hour}:00 with {hour_dist.max()} calls")
            self.insights.append(f"Most active day: {most_common_day}")

            current_step += 1
            self.update_progress.emit("Generating device usage insights...", int(current_step/total_steps*100))

            # Device usage insights
            top_devices = self.imei_counts.head(5)
            self.generate_bar_chart(
                top_devices,
                "Top 5 Most Used Devices",
                "IMEI",
                "Number of Calls",
                "device_usage.png"
            )
            self.visualizations['devices'] = "device_usage.png"
            self.insights.append(f"Top device IMEI: {top_devices.index[0]} used for {top_devices.iloc[0]} calls")

            current_step += 1
            self.update_progress.emit("Generating contact network analysis...", int(current_step/total_steps*100))

            # Contact network analysis
            top_contacts = pd.concat([outgoing_contacts, incoming_contacts]).nlargest(5)
            self.generate_bar_chart(
                top_contacts,
                "Top 5 Most Frequent Contacts",
                "Phone Number",
                "Number of Interactions",
                "top_contacts.png"
            )
            self.visualizations['contacts'] = "top_contacts.png"
            self.insights.append(f"Most frequent contact: {top_contacts.index[0]} with {top_contacts.iloc[0]} interactions")

            current_step += 1
            self.update_progress.emit("Finalizing results...", int(current_step/total_steps*100))

            # Last record information
            last_record = calls.iloc[-1]
            results.extend([
                ("Last Record Type", last_record["Call Type"]),
                ("Last Record Time", str(last_record["Event Date & Time"])),
                ("Last Location", f"Lat: {last_record['Latitude']:.6f}, Lon: {last_record['Longitude']:.6f}"),
                ("Last Cell Tower", last_record["Cell Details"])
            ])

            # Generate map
            self.generate_map(calls)
            self.visualizations['map'] = "AirtelTigo_CDR_Map.html"

            self.analysis_complete.emit(results, {
                'visualizations': self.visualizations,
                'insights': self.insights
            })

        except Exception as e:
            logger.error(f"Analysis error: {str(e)}")
            self.error_occurred.emit(f"Analysis failed: {str(e)}")

    def get_time_period(self, hour):
        if 0 <= hour < 6:
            return 'Night'
        elif 6 <= hour < 12:
            return 'Morning'
        elif 12 <= hour < 18:
            return 'Afternoon'
        else:
            return 'Evening'

    def generate_geolocation_plot(self, calls):
        plt.figure(figsize=(10, 6))
        calls = calls.dropna(subset=['Latitude', 'Longitude'])
        plt.scatter(calls["Longitude"], calls["Latitude"], alpha=0.6, edgecolors="w", s=100)
        plt.xlabel("Longitude")
        plt.ylabel("Latitude")
        plt.title("Call Locations")
        plt.grid(True)
        plt.savefig("geolocation_plot.png", dpi=100, bbox_inches='tight')
        plt.close()

    def generate_time_distribution_charts(self, calls):
        plt.figure(figsize=(10, 5))
        calls['Hour'].value_counts().sort_index().plot(kind='bar')
        plt.title("Call Distribution by Hour")
        plt.xlabel("Hour of Day")
        plt.ylabel("Number of Calls")
        plt.savefig("hourly_distribution.png", dpi=100, bbox_inches='tight')
        plt.close()
        
        plt.figure(figsize=(8, 5))
        calls['DayOfWeek'].value_counts().plot(kind='bar')
        plt.title("Call Distribution by Day of Week")
        plt.xlabel("Day of Week")
        plt.ylabel("Number of Calls")
        plt.savefig("day_distribution.png", dpi=100, bbox_inches='tight')
        plt.close()

    def generate_pie_chart(self, data, title, filename):
        plt.figure(figsize=(8, 8))
        data.plot(kind='pie', autopct='%1.1f%%', startangle=90)
        plt.title(title)
        plt.ylabel('')
        plt.savefig(filename, dpi=100, bbox_inches='tight')
        plt.close()

    def generate_histogram(self, data, title, xlabel, ylabel, filename, bins=10):
        plt.figure(figsize=(10, 6))
        plt.hist(data, bins=bins, edgecolor='black')
        plt.title(title)
        plt.xlabel(xlabel)
        plt.ylabel(ylabel)
        plt.grid(True)
        plt.savefig(filename, dpi=100, bbox_inches='tight')
        plt.close()

    def generate_bar_chart(self, data, title, xlabel, ylabel, filename):
        plt.figure(figsize=(10, 6))
        data.plot(kind='bar', color=ACCENT_COLOR)
        plt.title(title)
        plt.xlabel(xlabel)
        plt.ylabel(ylabel)
        plt.grid(True)
        plt.savefig(filename, dpi=100, bbox_inches='tight')
        plt.close()

    def generate_map(self, calls):
        try:
            calls = calls.dropna(subset=["Latitude", "Longitude"])
            if calls.empty:
                raise ValueError("No valid location data available")
            
            location_counts = calls.groupby(["Latitude", "Longitude"]).size().reset_index(name="count")
            highest_location = location_counts.loc[location_counts['count'].idxmax()]
            lat, lon = highest_location["Latitude"], highest_location["Longitude"]
            
            m = folium.Map(location=[lat, lon], zoom_start=12)
            
            for _, row in location_counts.iterrows():
                folium.Marker(
                    [row["Latitude"], row["Longitude"]],
                    popup=f"Cell Tower: {self.get_cell_tower_name(row['Latitude'], row['Longitude'], calls)}\nRecords: {row['count']}",
                    icon=folium.Icon(color='red', icon='tower-cell', prefix='fa')
                ).add_to(m)
            
            HeatMap(data=location_counts[["Latitude", "Longitude", "count"]], radius=15).add_to(m)
            
            m.save("AirtelTigo_CDR_Map.html")
            
        except Exception as e:
            logger.error(f"Map generation error: {str(e)}")
            raise

    def get_cell_tower_name(self, lat, lon, calls):
        tower = calls[(calls["Latitude"] == lat) & (calls["Longitude"] == lon)]["Cell Details"].mode()
        return tower[0] if not tower.empty else "Unknown Tower"

class AirtelTigoCDRAnalyzer(QMainWindow):
    def __init__(self, back_to_home_callback=None):
        super().__init__()
        self.back_to_home_callback = back_to_home_callback
        self.setWindowTitle("AirtelTigo CDR Analyzer")
        self.setGeometry(100, 100, 1200, 800)
        
        self.filename = None
        self.df = None
        self.map_path = os.path.join(os.getcwd(), "AirtelTigo_CDR_Map.html")
        self.doc_path = None
        self.analysis_thread = None
        self.visualizations = {}
        self.insights = []
        
        self.init_ui()
        self.create_home_page()
        self.setStyleSheet(self.get_stylesheet())
    
    def init_ui(self):
        self.central_widget = QWidget()
        self.setCentralWidget(self.central_widget)
        self.main_layout = QHBoxLayout(self.central_widget)
        self.main_layout.setContentsMargins(0, 0, 0, 0)
        self.main_layout.setSpacing(0)
        
        # Sidebar
        self.sidebar = QFrame()
        self.sidebar.setObjectName("sidebar")
        self.sidebar_width = 240
        self.sidebar.setFixedWidth(self.sidebar_width)
        self.sidebar_layout = QVBoxLayout(self.sidebar)
        self.sidebar_layout.setAlignment(Qt.AlignTop)
        self.sidebar_layout.setSpacing(10)
        self.sidebar_layout.setContentsMargins(0, 10, 0, 10)
        
        # Floating toggle button
        self.toggle_sidebar_btn = QPushButton("✕")
        self.toggle_sidebar_btn.setObjectName("toggleSidebar")
        self.toggle_sidebar_btn.setFixedSize(30, 30)
        self.toggle_sidebar_btn.clicked.connect(self.toggle_sidebar)
        self.toggle_sidebar_btn.setParent(self.central_widget)
        self.toggle_sidebar_btn.move(0, 10)
        self.toggle_sidebar_btn.raise_()
        
        # Sidebar content
        self.create_sidebar_buttons()
        self.sidebar_layout.addStretch()
        self.add_sidebar_logo()
        
        self.main_layout.addWidget(self.sidebar)
        
        # Content area
        self.content_stack = QStackedWidget()
        self.content_stack.setMinimumWidth(800)
        self.main_layout.addWidget(self.content_stack)
    
    def create_sidebar_buttons(self):
        buttons = [
            ("Load File", self.load_file, "standard"),
            ("Analyze Data", self.create_analysis_page, "standard"),
            ("Open Map", self.open_map, "standard"),
            ("Export Report", self.export_report, "standard"),
            ("Back to Home", self.back_to_home, "warning")
        ]
        
        for text, callback, btn_type in buttons:
            btn = QPushButton(text)
            if btn_type == "warning":
                btn.setProperty("class", "reset")
            btn.clicked.connect(callback)
            self.sidebar_layout.addWidget(btn)
    
    def add_sidebar_logo(self):
        logo_frame = QFrame()
        logo_frame.setObjectName("logo_frame")
        logo_layout = QVBoxLayout(logo_frame)
        logo_layout.setAlignment(Qt.AlignCenter)
        
        logo_label = QLabel()
        logo_path = os.path.join("assets", "airtel-tigo.png")
        if os.path.exists(logo_path):
            logo_pixmap = QPixmap(logo_path)
            if not logo_pixmap.isNull():
                logo_pixmap = logo_pixmap.scaled(100, 0, Qt.KeepAspectRatio, Qt.SmoothTransformation)
                logo_label.setPixmap(logo_pixmap)
                logo_label.setAlignment(Qt.AlignCenter)
        else:
            logo_label.setText("AirtelTigo")
            logo_label.setStyleSheet("font-size: 16px; font-weight: bold; color: white;")
        
        logo_layout.addWidget(logo_label)
        self.sidebar_layout.addWidget(logo_frame)
    
    def toggle_sidebar(self):
        if self.sidebar.width() > 0:
            self.animate_sidebar(0)
            self.toggle_sidebar_btn.setText("☰")
            self.toggle_sidebar_btn.setStyleSheet("""
                QPushButton {
                    background-color: #2C2C3A;
                    border-radius: 0 5px 5px 0;
                }
            """)
        else:
            self.animate_sidebar(self.sidebar_width)
            self.toggle_sidebar_btn.setText("✕")
            self.toggle_sidebar_btn.setStyleSheet("")
    
    def animate_sidebar(self, target_width):
        self.animation = QPropertyAnimation(self.sidebar, b"minimumWidth")
        self.animation.setDuration(200)
        self.animation.setStartValue(self.sidebar.width())
        self.animation.setEndValue(target_width)
        self.animation.setEasingCurve(QEasingCurve.InOutQuad)
        self.animation.start()
        
        btn_animation = QPropertyAnimation(self.toggle_sidebar_btn, b"pos")
        btn_animation.setDuration(200)
        btn_animation.setStartValue(self.toggle_sidebar_btn.pos())
        btn_animation.setEndValue(QPoint(target_width, 10))
        btn_animation.setEasingCurve(QEasingCurve.InOutQuad)
        btn_animation.start()
    
    def get_stylesheet(self):
        return f"""
            QMainWindow {{
                background-color: {BACKGROUND_COLOR};
                color: {TEXT_COLOR};
            }}
            
            QFrame#sidebar {{
                background-color: {SIDEBAR_COLOR};
                border: none;
            }}
            
            QPushButton#toggleSidebar {{
                background-color: transparent;
                color: white;
                font-size: 20px;
                font-weight: bold;
                border: none;
                padding: 5px;
            }}
            
            QPushButton#toggleSidebar:hover {{
                background-color: rgba(255, 255, 255, 0.1);
            }}
            
            QWidget#content_area {{
                background-color: {BACKGROUND_COLOR};
                padding: 20px;
            }}
            
            QTreeWidget {{
                background-color: {CONTENT_BG_COLOR};
                color: {TEXT_COLOR};
                border: 1px solid {BORDER_COLOR};
                font-size: 12px;
            }}
            
            QTreeWidget::item {{
                padding: 5px;
            }}
            
            QHeaderView::section {{
                background-color: {CONTENT_BG_COLOR};
                color: {TEXT_COLOR};
                padding: 5px;
                border: none;
            }}
            
            QPushButton {{
                background-color: {ACCENT_COLOR};
                color: {ACCENT_TEXT};
                font-size: 14px;
                font-weight: bold;
                padding: 8px 12px;
                border: none;
                border-radius: 5px;
                min-width: 100px;
                min-height: 40px;
            }}
            
            QPushButton:hover {{
                background-color: {HIGHLIGHT_COLOR};
            }}
            
            QPushButton.reset {{
                background-color: {ERROR_COLOR};
                color: white;
            }}
            
            QPushButton.reset:hover {{
                background-color: #C0392B;
            }}
            
            QProgressBar {{
                background: {CONTENT_BG_COLOR};
                border: 1px solid {BORDER_COLOR};
                text-align: center;
                color: {TEXT_COLOR};
            }}
            
            QProgressBar::chunk {{
                background-color: {ACCENT_COLOR};
            }}
            
            QLineEdit {{
                background: {CONTENT_BG_COLOR};
                border: 1px solid {BORDER_COLOR};
                padding: 5px;
                color: {TEXT_COLOR};
            }}
            
            QLabel {{
                color: {TEXT_COLOR};
            }}
            
            QLabel#file_label {{
                qproperty-alignment: AlignCenter;
                font-size: 14px;
            }}
            
            QLabel#file_label[error=true] {{
                color: {ERROR_COLOR};
            }}
            
            QLabel#file_label[success=true] {{
                color: {SUCCESS_COLOR};
            }}
            
            QScrollArea {{
                border: none;
            }}
            
            QTabWidget::pane {{
                border: 1px solid {BORDER_COLOR};
            }}
            
            QTabBar::tab {{
                background: {CONTENT_BG_COLOR};
                color: {TEXT_COLOR};
                padding: 8px;
                border: 1px solid {BORDER_COLOR};
                border-bottom: none;
            }}
            
            QTabBar::tab:selected {{
                background: {ACCENT_COLOR};
                color: {ACCENT_TEXT};
            }}
        """
    
    def create_home_page(self):
        if hasattr(self, 'home_page'):
            self.content_stack.setCurrentWidget(self.home_page)
            return
        
        self.home_page = QWidget()
        home_layout = QVBoxLayout(self.home_page)
        home_layout.setAlignment(Qt.AlignCenter)
        
        title = QLabel("AirtelTigo Call Detail Records Analyzer")
        title.setStyleSheet("font-size: 20px; font-weight: bold;")
        title.setAlignment(Qt.AlignCenter)
        home_layout.addWidget(title)
        
        self.file_label = QLabel("📂 No file selected")
        self.file_label.setObjectName("file_label")
        self.file_label.setProperty("error", True)
        self.file_label.style().polish(self.file_label)
        home_layout.addWidget(self.file_label)
        
        home_layout.addStretch()
        self.content_stack.addWidget(self.home_page)
        self.content_stack.setCurrentWidget(self.home_page)
    
    def create_analysis_page(self):
        if self.df is None:
            QMessageBox.warning(self, "Error", "Please load a file first!")
            return
        
        if hasattr(self, 'analysis_page'):
            self.content_stack.setCurrentWidget(self.analysis_page)
            return
        
        self.analysis_page = QWidget()
        analysis_layout = QVBoxLayout(self.analysis_page)
        
        title = QLabel("CDR Analysis Report")
        title.setStyleSheet("font-size: 18px; font-weight: bold;")
        title.setAlignment(Qt.AlignCenter)
        analysis_layout.addWidget(title)
        
        search_frame = QFrame()
        search_layout = QHBoxLayout(search_frame)
        
        self.search_input = QLineEdit()
        self.search_input.setPlaceholderText("Search...")
        self.search_input.textChanged.connect(self.filter_treeview)
        search_layout.addWidget(self.search_input)
        
        reset_btn = QPushButton("⭮ Reset")
        reset_btn.setProperty("class", "warning")
        reset_btn.clicked.connect(self.reset_filters)
        search_layout.addWidget(reset_btn)
        
        analysis_layout.addWidget(search_frame)
        
        self.tab_widget = QTabWidget()
        
        results_tab = QWidget()
        results_layout = QVBoxLayout(results_tab)
        
        self.tree = QTreeWidget()
        self.tree.setHeaderLabels(["Item", "Count"])
        self.tree.setColumnWidth(0, 300)
        self.tree.setColumnWidth(1, 500)
        results_layout.addWidget(self.tree)
        
        insights_label = QLabel("Key Insights")
        insights_label.setStyleSheet("font-size: 16px; font-weight: bold;")
        results_layout.addWidget(insights_label)
        
        self.insights_text = QLabel("Analysis insights will appear here...")
        self.insights_text.setWordWrap(True)
        results_layout.addWidget(self.insights_text)
        
        self.tab_widget.addTab(results_tab, "Results")
        
        viz_tab = QWidget()
        viz_layout = QVBoxLayout(viz_tab)
        
        self.viz_scroll = QScrollArea()
        self.viz_scroll.setWidgetResizable(True)
        self.viz_container = QWidget()
        self.viz_layout = QVBoxLayout(self.viz_container)
        
        self.viz_placeholder = QLabel("Visualizations will appear here after analysis")
        self.viz_placeholder.setAlignment(Qt.AlignCenter)
        self.viz_layout.addWidget(self.viz_placeholder)
        
        self.viz_scroll.setWidget(self.viz_container)
        viz_layout.addWidget(self.viz_scroll)
        
        self.tab_widget.addTab(viz_tab, "Visualizations")
        
        analysis_layout.addWidget(self.tab_widget)
        
        self.progress_bar = QProgressBar()
        self.progress_bar.setRange(0, 100)
        analysis_layout.addWidget(self.progress_bar)
        
        self.status_label = QLabel("Analyzing data... Please wait")
        self.status_label.setAlignment(Qt.AlignCenter)
        analysis_layout.addWidget(self.status_label)
        
        self.content_stack.addWidget(self.analysis_page)
        self.content_stack.setCurrentWidget(self.analysis_page)
        
        self.start_analysis()
    
    def start_analysis(self):
        if self.analysis_thread and self.analysis_thread.isRunning():
            self.analysis_thread.terminate()
        
        self.analysis_thread = AnalysisThread(self.df)
        self.analysis_thread.update_progress.connect(self.update_progress_status)
        self.analysis_thread.analysis_complete.connect(self.display_results)
        self.analysis_thread.error_occurred.connect(self.show_error)
        self.analysis_thread.start()
    
    def update_progress_status(self, message, percent):
        self.status_label.setText(message)
        self.progress_bar.setValue(percent)
    
    def display_results(self, results, analysis_data):
        self.tree.clear()
        self.visualizations = analysis_data['visualizations']
        self.insights = analysis_data['insights']
        
        for category, details in results:
            item = QTreeWidgetItem([category, details])
            self.tree.addTopLevelItem(item)
        
        self.insights_text.setText("\n• ".join([""] + self.insights))
        
        self.show_visualizations()
        
        self.status_label.setText("Analysis Complete")
    
    def show_visualizations(self):
        for i in reversed(range(self.viz_layout.count())): 
            self.viz_layout.itemAt(i).widget().setParent(None)
        
        if hasattr(self, 'viz_placeholder'):
            self.viz_layout.removeWidget(self.viz_placeholder)
            self.viz_placeholder.deleteLater()
            del self.viz_placeholder
        
        for viz_type, viz_path in self.visualizations.items():
            if viz_type == 'map':
                continue
            
            if os.path.exists(viz_path):
                title = QLabel(self.get_viz_title(viz_type))
                title.setStyleSheet("font-size: 14px; font-weight: bold;")
                self.viz_layout.addWidget(title)
                
                pixmap = QPixmap(viz_path)
                if not pixmap.isNull():
                    label = QLabel()
                    label.setPixmap(pixmap.scaled(800, 600, Qt.KeepAspectRatio, Qt.SmoothTransformation))
                    label.setAlignment(Qt.AlignCenter)
                    self.viz_layout.addWidget(label)
        
        self.viz_layout.addStretch()
    
    def get_viz_title(self, viz_type):
        titles = {
            'call_type': "Call Type Distribution",
            'duration': "Call Duration Distribution",
            'geolocation': "Call Locations",
            'hourly': "Hourly Call Distribution",
            'daily': "Daily Call Distribution",
            'devices': "Device Usage",
            'contacts': "Top Contacts"
        }
        return titles.get(viz_type, viz_type.replace('_', ' ').title())
    
    def filter_treeview(self):
        search_text = self.search_input.text().lower()
        
        for i in range(self.tree.topLevelItemCount()):
            item = self.tree.topLevelItem(i)
            category = item.text(0).lower()
            details = item.text(1).lower()
            item.setHidden(search_text not in category and search_text not in details)
    
    def reset_filters(self):
        self.search_input.clear()
        for i in range(self.tree.topLevelItemCount()):
            self.tree.topLevelItem(i).setHidden(False)
    
    def load_file(self):
        filename, _ = QFileDialog.getOpenFileName(
            self, "Open File", "", "Excel Files (*.xlsx *.xls);;CSV Files (*.csv)"
        )
        
        if filename:
            try:
                if filename.endswith('.csv'):
                    self.df = pd.read_csv(filename)
                else:
                    self.df = pd.read_excel(filename)

                self.df.columns = self.df.columns.str.strip()
                
                required_columns = ["Owner Number", "Outgoing", "Incoming", "Duration", 
                                  "Call Type", "Event Date & Time", "Latitude", "Longitude"]
                missing_columns = [col for col in required_columns if col not in self.df.columns]
                
                if missing_columns:
                    QMessageBox.critical(self, "Error", f"Missing required columns: {', '.join(missing_columns)}")
                    self.df = None
                    return

                self.filename = filename
                self.file_label.setText(f"📂 Loaded: {os.path.basename(filename)}")
                self.file_label.setProperty("error", False)
                self.file_label.setProperty("success", True)
                self.file_label.style().polish(self.file_label)
                QMessageBox.information(self, "Success", "File loaded successfully!")
                
            except Exception as e:
                logger.error(f"File loading error: {str(e)}")
                QMessageBox.critical(self, "Error", f"Failed to load file:\n{str(e)}")
                self.df = None
                self.file_label.setText("📂 Error loading file")
                self.file_label.setProperty("error", True)
                self.file_label.setProperty("success", False)
                self.file_label.style().polish(self.file_label)
    
    def export_report(self):
        if not hasattr(self, 'tree') or self.tree.topLevelItemCount() == 0:
            QMessageBox.warning(self, "Error", "No analysis results to export!")
            return
        
        filename, _ = QFileDialog.getSaveFileName(
            self, "Save Report", "", "Word Documents (*.docx)"
        )
        
        if not filename:
            return
        
        try:
            doc = Document()
            
            doc.add_heading('AirtelTigo CDR Analysis Report', 0)
            current_time = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            doc.add_paragraph(f"Report generated on: {current_time}")
            doc.add_paragraph(f"Source file: {os.path.basename(self.filename)}")
            doc.add_paragraph("\n")

            doc.add_heading('Summary Statistics', level=1)
            for i in range(self.tree.topLevelItemCount()):
                item = self.tree.topLevelItem(i)
                doc.add_paragraph(f"{item.text(0)}: {item.text(1)}")
            doc.add_paragraph("\n")

            doc.add_heading('Key Insights', level=1)
            for insight in self.insights:
                doc.add_paragraph(insight, style='List Bullet')
            doc.add_paragraph("\n")

            doc.add_heading('Visualizations', level=1)
            
            for viz_type in ['call_type', 'duration', 'hourly', 'daily', 'devices', 'contacts', 'geolocation']:
                if viz_type in self.visualizations and os.path.exists(self.visualizations[viz_type]):
                    doc.add_heading(self.get_viz_title(viz_type), level=2)
                    doc.add_picture(self.visualizations[viz_type], width=Inches(6))
                    doc.add_paragraph("\n")

            doc.save(filename)

            try:
                webbrowser.open(filename)
            except Exception as e:
                QMessageBox.warning(self, "Warning", f"Could not open document: {e}")

            QMessageBox.information(self, "Success", f"Report exported to {filename}")

        except Exception as e:
            logger.error(f"Report export error: {str(e)}")
            QMessageBox.critical(self, "Error", f"Failed to export report:\n{str(e)}")
    
    def open_map(self):
        if 'map' not in self.visualizations:
            QMessageBox.warning(self, "Error", "Map not generated yet. Please analyze data first.")
            return
        
        map_path = self.visualizations['map']
        if os.path.exists(map_path):
            webbrowser.open(map_path)
        else:
            QMessageBox.warning(self, "Error", "Map file not found. Please analyze data again.")
    
    def show_error(self, message):
        QMessageBox.critical(self, "Error", message)
        self.status_label.setText("Analysis Failed")
        self.progress_bar.setValue(0)
    
    def back_to_home(self):
        """Return to home with confirmation if analysis is in progress"""
        if hasattr(self, 'analysis_thread') and self.analysis_thread and self.analysis_thread.isRunning():
            reply = QMessageBox.question(
                self, 'Confirm Exit',
                'Analysis is in progress. Are you sure you want to exit?',
                QMessageBox.Yes | QMessageBox.No,
                QMessageBox.No
            )
            if reply == QMessageBox.No:
                return

        # Reset all data
        self.filename = None
        self.df = None
        self.map_path = os.path.join(os.getcwd(), "AirtelTigo_CDR_Map.html")
        self.doc_path = None
        self.visualizations = {}
        self.insights = []
        
        self.clear_content()
        
        self.file_label.setText("📂 No file selected")
        self.file_label.setProperty("error", True)
        self.file_label.style().polish(self.file_label)
        
        self.content_stack.setCurrentWidget(self.home_page)
        
        if self.back_to_home_callback:
            self.back_to_home_callback()
    
    def clear_content(self):
        if hasattr(self, 'analysis_page'):
            layout = self.analysis_page.layout()
            if layout:
                while layout.count():
                    item = layout.takeAt(0)
                    widget = item.widget()
                    if widget:
                        widget.deleteLater()
    
    def closeEvent(self, event):
        if hasattr(self, 'analysis_thread') and self.analysis_thread and self.analysis_thread.isRunning():
            self.analysis_thread.terminate()
        event.accept()

if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = AirtelTigoCDRAnalyzer()
    window.show()
    sys.exit(app.exec_())