import os
import sys
import pandas as pd
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout, 
    QPushButton, QLabel, QFrame, QFileDialog, QMessageBox, 
    QTreeWidget, QTreeWidgetItem, QProgressBar, QScrollArea, 
    QLineEdit, QSplitter, QStackedWidget, QTabWidget
)
from PyQt5.QtCore import Qt, QThread, pyqtSignal
from PyQt5.QtGui import QPixmap
from docx import Document
from docx.shared import Inches
import datetime
import subprocess
import logging

# Suppress libpng warning
os.environ['QT_LOGGING_RULES'] = '*.warning=false'

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('mobile_money_analyzer.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# THEME COLORS
BACKGROUND_COLOR = "#1E1E2F"  # Dark blue background
SIDEBAR_COLOR = "#1b2944"     # Sidebar background
CONTENT_BG_COLOR = "#2a3a5a"  # Content area background
TEXT_COLOR = "#FFFFFF"        # White text
BUTTON_COLOR = "#feca18"      # Yellow
BUTTON_TEXT_COLOR = "#0c0d41" # Dark blue
HIGHLIGHT_COLOR = "#e6b800"   # Slightly darker yellow on hover
ERROR_COLOR = "#FF6B6B"       # Light red for errors
SUCCESS_COLOR = "#4CAF50"     # Green for success states
BORDER_COLOR = "#3a4a6a"      # Border color

APP_STYLESHEET = f"""
    /* Main window */
    QMainWindow {{
        background-color: {BACKGROUND_COLOR};
        color: {TEXT_COLOR};
    }}
    
    /* Sidebar */
    QFrame#sidebar {{
        background-color: {SIDEBAR_COLOR};
        border-right: 1px solid {BORDER_COLOR};
    }}
    
    /* Content area */
    QWidget#content_area {{
        background-color: {BACKGROUND_COLOR};
        padding: 20px;
    }}
    
    /* Tree widget */
    QTreeWidget {{
        background-color: {CONTENT_BG_COLOR};
        color: {TEXT_COLOR};
        border: 1px solid {BORDER_COLOR};
        font-size: 12px;
    }}
    
    QTreeWidget::item {{
        padding: 5px;
    }}
    
    QHeaderView::section {{
        background-color: {CONTENT_BG_COLOR};
        color: {TEXT_COLOR};
        padding: 5px;
        border: none;
    }}
    
    /* Standard Buttons */
    QPushButton {{
        background-color: {BUTTON_COLOR};
        color: {BUTTON_TEXT_COLOR};
        font-size: 14px;
        font-weight: bold;
        padding: 8px 12px;
        border: none;
        border-radius: 5px;
        min-width: 100px;
        min-height: 40px;
    }}
    
    QPushButton:hover {{
        background-color: {HIGHLIGHT_COLOR};
    }}
    
    /* Sidebar Buttons */
    QFrame#sidebar QPushButton {{
        min-width: 180px;
        min-height: 45px;
    }}
    
    /* Important Buttons */
    QPushButton.important {{
        min-width: 150px;
        padding: 10px 15px;
    }}
    
    /* Progress bar */
    QProgressBar {{
        background: {CONTENT_BG_COLOR};
        border: 1px solid {BORDER_COLOR};
        text-align: center;
        color: {TEXT_COLOR};
    }}
    
    QProgressBar::chunk {{
        background-color: {BUTTON_COLOR};
    }}
    
    /* Line edits */
    QLineEdit {{
        background: {CONTENT_BG_COLOR};
        border: 1px solid {BORDER_COLOR};
        padding: 5px;
        color: {TEXT_COLOR};
    }}
    
    /* Labels */
    QLabel {{
        color: {TEXT_COLOR};
    }}
    
    QLabel#file_label {{
        qproperty-alignment: AlignCenter;
        font-size: 14px;
    }}
    
    QLabel#file_label[error=true] {{
        color: {ERROR_COLOR};
    }}
    
    QLabel#file_label[success=true] {{
        color: {SUCCESS_COLOR};
    }}
    
    /* Scroll areas */
    QScrollArea {{
        border: none;
    }}
    
    QScrollBar:vertical {{
        background: {CONTENT_BG_COLOR};
        width: 10px;
    }}
    
    QScrollBar::handle:vertical {{
        background: {BORDER_COLOR};
        min-height: 20px;
    }}
    
    /* Logo frame - matches sidebar background */
    QFrame#logo_frame {{
        background-color: {SIDEBAR_COLOR};
        padding: 5px;
        margin-top: 10px;
    }}
    
    /* Tab widget */
    QTabWidget::pane {{
        border: 1px solid {BORDER_COLOR};
    }}
    
    QTabBar::tab {{
        background: {CONTENT_BG_COLOR};
        color: {TEXT_COLOR};
        padding: 8px;
        border: 1px solid {BORDER_COLOR};
        border-bottom: none;
    }}
    
    QTabBar::tab:selected {{
        background: {BUTTON_COLOR};
        color: {BUTTON_TEXT_COLOR};
    }}
"""

class AnalysisThread(QThread):
    update_progress = pyqtSignal(str, int)
    analysis_complete = pyqtSignal(list, dict)
    error_occurred = pyqtSignal(str)

    def __init__(self, df):
        super().__init__()
        self.df = df
        self.visualizations = {}
        self.insights = []

    def run(self):
        try:
            if self.df is None:
                self.error_occurred.emit("No data loaded for analysis")
                return

            results = []
            total_steps = 10
            current_step = 0

            # Basic transaction analysis
            self.update_progress.emit("Analyzing transaction patterns...", int(current_step/total_steps*100))
            
            # 1. Total Transactions
            total_transactions = len(self.df)
            results.append(("Total Transactions", f"{total_transactions}"))
            
            # 2. Total Income
            total_income = self.df[self.df['TRANSACTION TYPE'] == 'CREDIT']['FROM AMOUNT'].sum()
            results.append(("Total Income", f"₵{total_income:,.2f}"))
            
            # 3. Total Expenses
            total_expenses = self.df[self.df['TRANSACTION TYPE'] == 'DEBIT']['FROM AMOUNT'].sum()
            results.append(("Total Expenses", f"₵{total_expenses:,.2f}"))
            
            # 4. Net Balance
            net_balance = total_income - total_expenses
            results.append(("Net Balance", f"₵{net_balance:,.2f}"))
            
            # 5. Average Transaction Amount
            avg_transaction = self.df['FROM AMOUNT'].mean()
            results.append(("Average Transaction Amount", f"₵{avg_transaction:,.2f}"))
            
            current_step += 1
            self.update_progress.emit("Analyzing transaction types...", int(current_step/total_steps*100))

            # 6. Transaction Types
            transaction_counts = self.df['TRANSACTION TYPE'].value_counts()
            results.append(("Transaction Types", "Count"))
            
            for trans_type, count in transaction_counts.items():
                results.append((f"{trans_type}", f"{count}"))
                
            # Generate transaction type pie chart
            self.generate_pie_chart(
                transaction_counts,
                "Transaction Type Distribution",
                "transaction_types.png"
            )
            self.visualizations['transaction_types'] = "transaction_types.png"
            self.insights.append(f"Transaction types: {', '.join([f'{k} ({v})' for k,v in transaction_counts.items()])}")

            current_step += 1
            self.update_progress.emit("Analyzing top senders...", int(current_step/total_steps*100))

            # 7. Top Senders
            top_senders = self.df.groupby(['FROM ACCOUNT', 'FROM ACCOUNT NAME', 'FROM PHONE NUMBER']).size().reset_index(name='COUNT')
            top_senders = top_senders.sort_values(by='COUNT', ascending=False).head(5)
            results.append(("Top Senders", "Count"))
            
            for i, (_, row) in enumerate(top_senders.iterrows(), 1):
                results.append((f"{row['FROM ACCOUNT NAME']} ({row['FROM PHONE NUMBER']})", f"{row['COUNT']}"))
                
            # Generate top senders bar chart
            self.generate_bar_chart(
                top_senders.set_index('FROM ACCOUNT NAME')['COUNT'],
                "Top 5 Senders",
                "Sender",
                "Number of Transactions",
                "top_senders.png"
            )
            self.visualizations['top_senders'] = "top_senders.png"
            self.insights.append(f"Top sender: {top_senders.iloc[0]['FROM ACCOUNT NAME']} with {top_senders.iloc[0]['COUNT']} transactions")

            current_step += 1
            self.update_progress.emit("Analyzing top receivers...", int(current_step/total_steps*100))

            # 8. Top Receivers
            top_receivers = self.df.groupby(['TO ACCOUNT', 'TO ACCOUNT NAME', 'TO PHONE NUMBER']).size().reset_index(name='COUNT')
            top_receivers = top_receivers.sort_values(by='COUNT', ascending=False).head(5)
            results.append(("Top Receivers", "Count"))
            
            for i, (_, row) in enumerate(top_receivers.iterrows(), 1):
                results.append((f"{row['TO ACCOUNT NAME']} ({row['TO PHONE NUMBER']})", f"{row['COUNT']}"))
                
            # Generate top receivers bar chart
            self.generate_bar_chart(
                top_receivers.set_index('TO ACCOUNT NAME')['COUNT'],
                "Top 5 Receivers",
                "Receiver",
                "Number of Transactions",
                "top_receivers.png"
            )
            self.visualizations['top_receivers'] = "top_receivers.png"
            self.insights.append(f"Top receiver: {top_receivers.iloc[0]['TO ACCOUNT NAME']} with {top_receivers.iloc[0]['COUNT']} transactions")

            current_step += 1
            self.update_progress.emit("Analyzing transaction amounts...", int(current_step/total_steps*100))

            # 9. Transaction Amount Distribution
            self.generate_histogram(
                self.df['FROM AMOUNT'],
                "Transaction Amount Distribution",
                "Amount (₵)",
                "Number of Transactions",
                "amount_distribution.png",
                bins=20
            )
            self.visualizations['amount_dist'] = "amount_distribution.png"
            self.insights.append(f"Average transaction amount: ₵{avg_transaction:,.2f}")

            current_step += 1
            self.update_progress.emit("Generating time-based analysis...", int(current_step/total_steps*100))

            # 10. Time-based analysis (if date column exists)
            if 'DATE' in self.df.columns:
                try:
                    self.df['DATE'] = pd.to_datetime(self.df['DATE'])
                    self.df['HOUR'] = self.df['DATE'].dt.hour
                    self.df['DAY_OF_WEEK'] = self.df['DATE'].dt.day_name()
                    
                    # Hourly distribution
                    hourly_counts = self.df['HOUR'].value_counts().sort_index()
                    self.generate_bar_chart(
                        hourly_counts,
                        "Hourly Transaction Distribution",
                        "Hour of Day",
                        "Number of Transactions",
                        "hourly_distribution.png"
                    )
                    self.visualizations['hourly'] = "hourly_distribution.png"
                    
                    peak_hour = hourly_counts.idxmax()
                    self.insights.append(f"Peak transaction hour: {peak_hour}:00 with {hourly_counts.max()} transactions")
                    
                    # Day of week distribution
                    day_counts = self.df['DAY_OF_WEEK'].value_counts()
                    self.generate_bar_chart(
                        day_counts,
                        "Day of Week Distribution",
                        "Day",
                        "Number of Transactions",
                        "day_distribution.png"
                    )
                    self.visualizations['daily'] = "day_distribution.png"
                    
                    peak_day = day_counts.idxmax()
                    self.insights.append(f"Most active day: {peak_day} with {day_counts.max()} transactions")
                except:
                    pass

            current_step += 1
            self.update_progress.emit("Finalizing results...", int(current_step/total_steps*100))

            self.analysis_complete.emit(results, {
                'visualizations': self.visualizations,
                'insights': self.insights
            })

        except Exception as e:
            logger.error(f"Analysis error: {str(e)}")
            self.error_occurred.emit(f"Analysis failed: {str(e)}")

    def generate_pie_chart(self, data, title, filename):
        """Generate a pie chart visualization"""
        plt.figure(figsize=(8, 8))
        data.plot(kind='pie', autopct='%1.1f%%', startangle=90)
        plt.title(title)
        plt.ylabel('')
        plt.savefig(filename, dpi=100, bbox_inches='tight')
        plt.close()

    def generate_histogram(self, data, title, xlabel, ylabel, filename, bins=10):
        """Generate a histogram visualization"""
        plt.figure(figsize=(10, 6))
        plt.hist(data, bins=bins, edgecolor='black')
        plt.title(title)
        plt.xlabel(xlabel)
        plt.ylabel(ylabel)
        plt.grid(True)
        plt.savefig(filename, dpi=100, bbox_inches='tight')
        plt.close()

    def generate_bar_chart(self, data, title, xlabel, ylabel, filename):
        """Generate a bar chart visualization"""
        plt.figure(figsize=(10, 6))
        data.plot(kind='bar')
        plt.title(title)
        plt.xlabel(xlabel)
        plt.ylabel(ylabel)
        plt.grid(True)
        plt.savefig(filename, dpi=100, bbox_inches='tight')
        plt.close()


class MobileMoneyAnalyzer(QMainWindow):
    def __init__(self, back_to_home_callback):
        super().__init__()
        self.back_to_home_callback = back_to_home_callback
        self.setWindowTitle("MTN Mobile Money Transaction Analyzer")
        self.setGeometry(100, 100, 1200, 800)
        self.setStyleSheet(APP_STYLESHEET)
        
        # Initialize variables
        self.filename = None
        self.df = None
        self.doc_path = None
        self.analysis_thread = None
        self.visualizations = {}
        self.insights = []
        
        # Create main widgets
        self.init_ui()
        self.create_home_page()
    
    def init_ui(self):
        """Initialize main UI components"""
        self.central_widget = QWidget()
        self.setCentralWidget(self.central_widget)
        self.main_layout = QHBoxLayout(self.central_widget)
        self.main_layout.setContentsMargins(0, 0, 0, 0)
        self.main_layout.setSpacing(0)
        
        # Create sidebar with styled buttons
        self.sidebar = QFrame()
        self.sidebar.setObjectName("sidebar")
        self.sidebar.setFixedWidth(240)
        self.sidebar_layout = QVBoxLayout(self.sidebar)
        self.sidebar_layout.setAlignment(Qt.AlignTop)
        self.sidebar_layout.setSpacing(10)
        self.sidebar_layout.setContentsMargins(10, 20, 10, 10)
        
        # Add buttons to sidebar
        self.create_sidebar_buttons()
        self.sidebar_layout.addStretch()
        self.add_sidebar_logo()
        
        self.main_layout.addWidget(self.sidebar)
        
        # Create content area
        self.content_stack = QStackedWidget()
        self.main_layout.addWidget(self.content_stack)
    
    def create_sidebar_buttons(self):
        """Create sidebar navigation buttons"""
        buttons = [
            ("Load File", self.load_file, "standard"),
            ("Analyze Data", self.create_analysis_page, "standard"),
            ("Export Report", self.export_report, "standard"),
            ("Back to Home", self.back_to_home, "warning")
        ]
        
        for text, callback, btn_type in buttons:
            btn = QPushButton(text)
            btn.setStyleSheet(f"""
                QPushButton {{
                    background-color: {BUTTON_COLOR};
                    color: {BUTTON_TEXT_COLOR};
                    font-size: 14px;
                    font-weight: bold;
                    padding: 10px;
                    border: none;
                    border-radius: 5px;
                    min-width: 180px;
                    min-height: 45px;
                }}
                QPushButton:hover {{
                    background-color: {HIGHLIGHT_COLOR};
                }}
                QPushButton:disabled {{
                    background-color: #7a7a7a;
                    color: #cccccc;
                }}
            """)
            btn.clicked.connect(callback)
            self.sidebar_layout.addWidget(btn)
            
            if text == "Export Report":
                self.export_button = btn
                self.export_button.setEnabled(False)
    
    def add_sidebar_logo(self):
        """Add Mobile Money logo to sidebar with matching background"""
        logo_frame = QFrame()
        logo_frame.setObjectName("logo_frame")
        logo_frame.setStyleSheet(f"background-color: {SIDEBAR_COLOR}; border: none;")
        logo_layout = QVBoxLayout(logo_frame)
        logo_layout.setAlignment(Qt.AlignCenter)
        logo_layout.setContentsMargins(5, 5, 5, 5)
        
        logo_label = QLabel()
        logo_path = os.path.join("assets", "momo.png")
        if os.path.exists(logo_path):
            logo_pixmap = QPixmap(logo_path)
            if not logo_pixmap.isNull():
                # Scale logo to appropriate size
                logo_pixmap = logo_pixmap.scaled(100, 100, Qt.KeepAspectRatio, Qt.SmoothTransformation)
                logo_label.setPixmap(logo_pixmap)
                logo_label.setAlignment(Qt.AlignCenter)
        else:
            # Fallback if logo not found
            logo_label.setText("Mobile Money")
            logo_label.setStyleSheet("""
                font-size: 16px; 
                font-weight: bold; 
                color: white;
            """)
        
        logo_layout.addWidget(logo_label)
        self.sidebar_layout.addWidget(logo_frame)
    
    def back_to_home(self):
        """Reset the app and return to home screen."""
        # Reset all variables
        self.filename = None
        self.df = None
        self.doc_path = None
        self.visualizations = {}
        self.insights = []
        
        # Clear the content
        self.clear_content()
        
        # Reset file label
        self.file_label.setText("📂 No file selected")
        self.file_label.setProperty("error", True)
        self.file_label.style().polish(self.file_label)
        
        # Disable export button
        self.export_button.setEnabled(False)
        
        # Return to home page
        self.content_stack.setCurrentWidget(self.home_page)
        
        # If we have a callback, use it to exit completely
        if self.back_to_home_callback:
            self.back_to_home_callback()
    
    def clear_content(self):
        """Clear the content area"""
        if hasattr(self, 'analysis_page'):
            layout = self.analysis_page.layout()
            if layout:
                while layout.count():
                    item = layout.takeAt(0)
                    widget = item.widget()
                    if widget:
                        widget.deleteLater()
    
    def create_home_page(self):
        """Create the home page"""
        if hasattr(self, 'home_page'):
            self.content_stack.setCurrentWidget(self.home_page)
            return
        
        self.home_page = QWidget()
        home_layout = QVBoxLayout(self.home_page)
        home_layout.setAlignment(Qt.AlignCenter)
        
        title = QLabel("MTN Mobile Money Transaction Analyzer")
        title.setStyleSheet("font-size: 20px; font-weight: bold;")
        title.setAlignment(Qt.AlignCenter)
        home_layout.addWidget(title)
        
        self.file_label = QLabel("📂 No file selected")
        self.file_label.setObjectName("file_label")
        self.file_label.setProperty("error", True)
        self.file_label.style().polish(self.file_label)
        home_layout.addWidget(self.file_label)
        
        home_layout.addStretch()
        self.content_stack.addWidget(self.home_page)
        self.content_stack.setCurrentWidget(self.home_page)
    
    def format_ghs_amount(self, amount):
        """Format amount with Ghana Cedi symbol"""
        try:
            return f"Gh₵{float(amount):,.2f}"
        except (ValueError, TypeError):
            return str(amount)
    
    def create_analysis_page(self):
        """Create the analysis page with results"""
        if self.df is None:
            QMessageBox.warning(self, "Error", "Please load a file first!")
            return
        
        if hasattr(self, 'analysis_page'):
            self.content_stack.setCurrentWidget(self.analysis_page)
            return
        
        self.analysis_page = QWidget()
        analysis_layout = QVBoxLayout(self.analysis_page)
        
        title = QLabel("Transaction Analysis Report")
        title.setStyleSheet("font-size: 18px; font-weight: bold;")
        title.setAlignment(Qt.AlignCenter)
        analysis_layout.addWidget(title)
        
        search_frame = QFrame()
        search_layout = QHBoxLayout(search_frame)
        
        self.search_input = QLineEdit()
        self.search_input.setPlaceholderText("Search...")
        self.search_input.textChanged.connect(self.filter_treeview)
        search_layout.addWidget(self.search_input)
        
        reset_btn = QPushButton("⭮ Reset")
        reset_btn.setStyleSheet(f"""
            QPushButton {{
                background-color: {ERROR_COLOR};
                color: white;
                font-size: 14px;
                font-weight: bold;
                padding: 8px 12px;
                border: none;
                border-radius: 5px;
            }}
            QPushButton:hover {{
                background-color: #C0392B;
            }}
        """)
        reset_btn.clicked.connect(self.reset_filters)
        search_layout.addWidget(reset_btn)
        
        analysis_layout.addWidget(search_frame)
        
        # Create tab widget for results and visualizations
        self.tab_widget = QTabWidget()
        
        # Results tab
        results_tab = QWidget()
        results_layout = QVBoxLayout(results_tab)
        
        self.tree = QTreeWidget()
        self.tree.setHeaderLabels(["Metric", "Value"])
        self.tree.setColumnWidth(0, 300)
        self.tree.setColumnWidth(1, 500)
        results_layout.addWidget(self.tree)
        
        # Insights section
        insights_label = QLabel("Key Insights")
        insights_label.setStyleSheet("font-size: 16px; font-weight: bold;")
        results_layout.addWidget(insights_label)
        
        self.insights_text = QLabel("Analysis insights will appear here...")
        self.insights_text.setWordWrap(True)
        results_layout.addWidget(self.insights_text)
        
        self.tab_widget.addTab(results_tab, "Results")
        
        # Visualizations tab
        viz_tab = QWidget()
        viz_layout = QVBoxLayout(viz_tab)
        
        self.viz_scroll = QScrollArea()
        self.viz_scroll.setWidgetResizable(True)
        self.viz_container = QWidget()
        self.viz_layout = QVBoxLayout(self.viz_container)
        
        # Placeholder for visualizations
        self.viz_placeholder = QLabel("Visualizations will appear here after analysis")
        self.viz_placeholder.setAlignment(Qt.AlignCenter)
        self.viz_layout.addWidget(self.viz_placeholder)
        
        self.viz_scroll.setWidget(self.viz_container)
        viz_layout.addWidget(self.viz_scroll)
        
        self.tab_widget.addTab(viz_tab, "Visualizations")
        
        # Enhanced Transactions tab
        transactions_tab = QWidget()
        transactions_layout = QVBoxLayout(transactions_tab)
        
        # Transaction filter controls
        filter_frame = QFrame()
        filter_layout = QHBoxLayout(filter_frame)
        
        # Type filter
        type_label = QLabel("Type:")
        self.type_filter = QLineEdit()
        self.type_filter.setPlaceholderText("CREDIT/DEBIT")
        filter_layout.addWidget(type_label)
        filter_layout.addWidget(self.type_filter)
        
        # From filter
        from_label = QLabel("From:")
        self.from_filter = QLineEdit()
        self.from_filter.setPlaceholderText("Name/Phone")
        filter_layout.addWidget(from_label)
        filter_layout.addWidget(self.from_filter)
        
        # To filter
        to_label = QLabel("To:")
        self.to_filter = QLineEdit()
        self.to_filter.setPlaceholderText("Name/Phone")
        filter_layout.addWidget(to_label)
        filter_layout.addWidget(self.to_filter)
        
        # Amount range
        amount_label = QLabel("Amount:")
        self.min_amount = QLineEdit()
        self.min_amount.setPlaceholderText("Min")
        self.min_amount.setFixedWidth(80)
        self.max_amount = QLineEdit()
        self.max_amount.setPlaceholderText("Max")
        self.max_amount.setFixedWidth(80)
        filter_layout.addWidget(amount_label)
        filter_layout.addWidget(self.min_amount)
        filter_layout.addWidget(self.max_amount)
        
        # Message filter
        msg_label = QLabel("Message:")
        self.msg_filter = QLineEdit()
        self.msg_filter.setPlaceholderText("Contains text")
        filter_layout.addWidget(msg_label)
        filter_layout.addWidget(self.msg_filter)
        
        # Apply filter button
        apply_filter_btn = QPushButton("Apply Filters")
        apply_filter_btn.clicked.connect(self.filter_transactions)
        filter_layout.addWidget(apply_filter_btn)
        
        # Clear filter button
        clear_filter_btn = QPushButton("Clear")
        clear_filter_btn.clicked.connect(self.clear_transaction_filters)
        filter_layout.addWidget(clear_filter_btn)
        
        transactions_layout.addWidget(filter_frame)
        
        # Create transaction tree widget with enhanced columns
        self.transactions_tree = QTreeWidget()
        self.transactions_tree.setHeaderLabels([
            "Type", 
            "From (Name - Phone)", 
            "Amount (₵)", 
            "To (Name - Phone)", 
            "Date", 
            "Reference",
            "Message"
        ])
        self.transactions_tree.setColumnWidth(0, 100)  # Type
        self.transactions_tree.setColumnWidth(1, 200)  # From
        self.transactions_tree.setColumnWidth(2, 100)  # Amount
        self.transactions_tree.setColumnWidth(3, 200)  # To
        self.transactions_tree.setColumnWidth(4, 120)  # Date
        self.transactions_tree.setColumnWidth(5, 150)  # Reference
        self.transactions_tree.setColumnWidth(6, 250)  # Message
        
        # Enable sorting
        self.transactions_tree.setSortingEnabled(True)
        
        # Add scroll area
        transactions_scroll = QScrollArea()
        transactions_scroll.setWidgetResizable(True)
        transactions_scroll.setWidget(self.transactions_tree)
        transactions_layout.addWidget(transactions_scroll)
        
        self.tab_widget.addTab(transactions_tab, "All Transactions")
        
        analysis_layout.addWidget(self.tab_widget)
        
        self.progress_bar = QProgressBar()
        self.progress_bar.setRange(0, 100)
        analysis_layout.addWidget(self.progress_bar)
        
        self.status_label = QLabel("Analyzing data... Please wait")
        self.status_label.setAlignment(Qt.AlignCenter)
        analysis_layout.addWidget(self.status_label)
        
        self.content_stack.addWidget(self.analysis_page)
        self.content_stack.setCurrentWidget(self.analysis_page)
        
        self.start_analysis()
    
    def populate_transactions(self):
        """Populate the transactions tree with all transactions (enhanced version)"""
        self.transactions_tree.clear()
        
        # Check required columns
        required_cols = [
            'TRANSACTION TYPE', 
            'FROM ACCOUNT NAME', 
            'FROM PHONE NUMBER', 
            'FROM AMOUNT',
            'TO ACCOUNT NAME',
            'TO PHONE NUMBER'
        ]
        
        if not all(col in self.df.columns for col in required_cols):
            QMessageBox.warning(self, "Warning", "Some required columns are missing in the data")
            return
        
        # Check for optional columns
        has_date = 'DATE' in self.df.columns
        has_reference = 'TRANSACTION REFERENCE' in self.df.columns
        has_message = 'MESSAGE' in self.df.columns
        
        for _, row in self.df.iterrows():
            # Create from and to account strings
            from_account = f"{row['FROM ACCOUNT NAME']} - {row['FROM PHONE NUMBER']}"
            to_account = f"{row['TO ACCOUNT NAME']} - {row['TO PHONE NUMBER']}"
            
            # Create item with main columns
            item = QTreeWidgetItem([
                row['TRANSACTION TYPE'],
                from_account,
                self.format_ghs_amount(row['FROM AMOUNT']),
                to_account
            ])
            
            # Add date if available
            if has_date:
                date_str = str(row['DATE'])
                item.setText(4, date_str)
            
            # Add reference if available
            if has_reference:
                item.setText(5, str(row['TRANSACTION REFERENCE']))
            
            # Add message if available
            if has_message:
                item.setText(6, str(row['MESSAGE']))
            
            self.transactions_tree.addTopLevelItem(item)
    
    def filter_transactions(self):
        """Filter transactions based on user input"""
        type_filter = self.type_filter.text().upper()
        from_filter = self.from_filter.text().lower()
        to_filter = self.to_filter.text().lower()
        msg_filter = self.msg_filter.text().lower()
        
        try:
            min_amount = float(self.min_amount.text()) if self.min_amount.text() else None
            max_amount = float(self.max_amount.text()) if self.max_amount.text() else None
        except ValueError:
            QMessageBox.warning(self, "Invalid Input", "Please enter valid numbers for amount range")
            return
        
        for i in range(self.transactions_tree.topLevelItemCount()):
            item = self.transactions_tree.topLevelItem(i)
            item_type = item.text(0).upper()
            from_text = item.text(1).lower()
            to_text = item.text(3).lower()
            msg_text = item.text(6).lower() if item.text(6) else ""
            
            try:
                amount = float(item.text(2).replace('₵', '').replace(',', ''))
            except ValueError:
                amount = 0
            
            # Apply filters
            type_match = not type_filter or type_filter in item_type
            from_match = not from_filter or from_filter in from_text
            to_match = not to_filter or to_filter in to_text
            msg_match = not msg_filter or msg_filter in msg_text
            min_amount_match = min_amount is None or amount >= min_amount
            max_amount_match = max_amount is None or amount <= max_amount
            
            item.setHidden(not (type_match and from_match and to_match and 
                              msg_match and min_amount_match and max_amount_match))
    
    def clear_transaction_filters(self):
        """Clear all transaction filters"""
        self.type_filter.clear()
        self.from_filter.clear()
        self.to_filter.clear()
        self.msg_filter.clear()
        self.min_amount.clear()
        self.max_amount.clear()
        
        for i in range(self.transactions_tree.topLevelItemCount()):
            self.transactions_tree.topLevelItem(i).setHidden(False)
    
    def start_analysis(self):
        """Start the analysis in a background thread"""
        if self.analysis_thread and self.analysis_thread.isRunning():
            self.analysis_thread.terminate()
        
        self.analysis_thread = AnalysisThread(self.df)
        self.analysis_thread.update_progress.connect(self.update_progress_status)
        self.analysis_thread.analysis_complete.connect(self.display_results)
        self.analysis_thread.error_occurred.connect(self.show_error)
        self.analysis_thread.start()
    
    def update_progress_status(self, message, percent):
        """Update progress status"""
        self.status_label.setText(message)
        self.progress_bar.setValue(percent)
    
    def display_results(self, results, analysis_data):
        """Display analysis results and visualizations"""
        self.tree.clear()
        self.visualizations = analysis_data['visualizations']
        self.insights = analysis_data['insights']
        
        # Display results in tree widget
        for category, details in results:
            item = QTreeWidgetItem([str(category), str(details)])
            self.tree.addTopLevelItem(item)
        
        # Display insights
        self.insights_text.setText("\n• ".join([""] + self.insights))
        
        # Display visualizations
        self.show_visualizations()
        
        # Populate transactions
        self.populate_transactions()
        
        # Enable export button
        self.export_button.setEnabled(True)
        
        self.status_label.setText("Analysis Complete")
    
    def show_visualizations(self):
        """Display all generated visualizations"""
        # Clear previous visualizations
        for i in reversed(range(self.viz_layout.count())): 
            self.viz_layout.itemAt(i).widget().setParent(None)
        
        # Remove placeholder
        if hasattr(self, 'viz_placeholder'):
            self.viz_layout.removeWidget(self.viz_placeholder)
            self.viz_placeholder.deleteLater()
            del self.viz_placeholder
        
        # Add each visualization
        for viz_type, viz_path in self.visualizations.items():
            if os.path.exists(viz_path):
                # Add title
                title = QLabel(self.get_viz_title(viz_type))
                title.setStyleSheet("font-size: 14px; font-weight: bold;")
                self.viz_layout.addWidget(title)
                
                # Add image
                pixmap = QPixmap(viz_path)
                if not pixmap.isNull():
                    label = QLabel()
                    label.setPixmap(pixmap.scaled(800, 600, Qt.KeepAspectRatio, Qt.SmoothTransformation))
                    label.setAlignment(Qt.AlignCenter)
                    self.viz_layout.addWidget(label)
        
        self.viz_layout.addStretch()
    
    def get_viz_title(self, viz_type):
        """Get display title for visualization type"""
        titles = {
            'transaction_types': "Transaction Type Distribution",
            'top_senders': "Top 5 Senders",
            'top_receivers': "Top 5 Receivers",
            'amount_dist': "Transaction Amount Distribution",
            'hourly': "Hourly Transaction Distribution",
            'daily': "Daily Transaction Distribution"
        }
        return titles.get(viz_type, viz_type.replace('_', ' ').title())
    
    def filter_treeview(self):
        """Filter tree items based on search text"""
        search_text = self.search_input.text().lower()
        
        for i in range(self.tree.topLevelItemCount()):
            item = self.tree.topLevelItem(i)
            category = item.text(0).lower()
            details = item.text(1).lower()
            item.setHidden(search_text not in category and search_text not in details)
    
    def reset_filters(self):
        """Reset all filters"""
        self.search_input.clear()
        for i in range(self.tree.topLevelItemCount()):
            self.tree.topLevelItem(i).setHidden(False)
    
    def load_file(self):
        """Load Excel file"""
        filename, _ = QFileDialog.getOpenFileName(
            self, "Open File", "", "Excel Files (*.xlsx *.xls);;CSV Files (*.csv)"
        )
        
        if filename:
            try:
                if filename.endswith('.csv'):
                    self.df = pd.read_csv(filename)
                else:
                    self.df = pd.read_excel(filename, engine='openpyxl')

                # Normalize column names
                self.df.columns = [col.strip().upper() for col in self.df.columns]

                # Check for required columns
                required_columns = ["TRANSACTION TYPE", "FROM AMOUNT"]
                missing_columns = [col for col in required_columns if col not in self.df.columns]
                
                if missing_columns:
                    QMessageBox.critical(self, "Error", f"Missing required columns: {', '.join(missing_columns)}")
                    self.df = None
                    return

                self.filename = filename
                self.file_label.setText(f"📂 Loaded: {os.path.basename(filename)}")
                self.file_label.setProperty("error", False)
                self.file_label.setProperty("success", True)
                self.file_label.style().polish(self.file_label)
                
                # Disable export button until analysis is complete
                self.export_button.setEnabled(False)
                
                QMessageBox.information(self, "Success", "File loaded successfully!")
                
            except Exception as e:
                logger.error(f"File loading error: {str(e)}")
                QMessageBox.critical(self, "Error", f"Failed to load file:\n{str(e)}")
                self.df = None
                self.file_label.setText("📂 Error loading file")
                self.file_label.setProperty("error", True)
                self.file_label.setProperty("success", False)
                self.file_label.style().polish(self.file_label)
    
    def export_report(self):
        """Export report to Word document with all enhancements"""
        if not hasattr(self, 'tree') or self.tree.topLevelItemCount() == 0:
            QMessageBox.warning(self, "Error", "No analysis results to export!")
            return
        
        filename, _ = QFileDialog.getSaveFileName(
            self, "Save Report", "", "Word Documents (*.docx)"
        )
        
        if not filename:
            return
        
        try:
            doc = Document()
            
            # Add title and metadata
            doc.add_heading('MTN Mobile Money Transaction Analysis Report', 0)
            current_time = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            doc.add_paragraph(f"Report generated on: {current_time}")
            doc.add_paragraph(f"Source file: {os.path.basename(self.filename)}")
            doc.add_paragraph("\n")

            # Add summary statistics
            doc.add_heading('Summary Statistics', level=1)
            for i in range(self.tree.topLevelItemCount()):
                item = self.tree.topLevelItem(i)
                doc.add_paragraph(f"{item.text(0)}: {item.text(1)}")
            doc.add_paragraph("\n")

            # Add insights
            doc.add_heading('Key Insights', level=1)
            for insight in self.insights:
                doc.add_paragraph(insight, style='List Bullet')
            doc.add_paragraph("\n")

            # Add enhanced transactions section
            doc.add_heading('Transaction Details', level=1)
            
            # Create table for transactions
            table = doc.add_table(rows=1, cols=7)
            table.style = 'Table Grid'
            
            # Header row
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Type'
            hdr_cells[1].text = 'From Account'
            hdr_cells[2].text = 'Amount'
            hdr_cells[3].text = 'To Account'
            hdr_cells[4].text = 'Date'
            hdr_cells[5].text = 'Reference'
            hdr_cells[6].text = 'Message'
            
            # Add transaction rows
            for _, row in self.df.iterrows():
                row_cells = table.add_row().cells
                row_cells[0].text = str(row['TRANSACTION TYPE'])
                from_account = f"{row['FROM ACCOUNT NAME']} - {row['FROM PHONE NUMBER']}"
                row_cells[1].text = from_account
                row_cells[2].text = self.format_ghs_amount(row['FROM AMOUNT'])
                to_account = f"{row['TO ACCOUNT NAME']} - {row['TO PHONE NUMBER']}"
                row_cells[3].text = to_account
                row_cells[4].text = str(row.get('DATE', ''))
                row_cells[5].text = str(row.get('TRANSACTION REFERENCE', ''))
                row_cells[6].text = str(row.get('MESSAGE', ''))
            
            doc.add_paragraph("\n")

            # Add visualizations
            doc.add_heading('Visualizations', level=1)
            
            # Add each visualization that exists in both UI and filesystem
            for viz_type in ['transaction_types', 'top_senders', 'top_receivers', 'amount_dist', 'hourly', 'daily']:
                if viz_type in self.visualizations and os.path.exists(self.visualizations[viz_type]):
                    doc.add_heading(self.get_viz_title(viz_type), level=2)
                    doc.add_picture(self.visualizations[viz_type], width=Inches(6))
                    doc.add_paragraph("\n")

            # Save document
            doc.save(filename)

            try:
                os.startfile(filename)  # For Windows
            except:
                subprocess.run(["open", filename])  # For macOS
                subprocess.run(["xdg-open", filename])  # For Linux

            QMessageBox.information(self, "Success", f"Report exported to {filename}")

        except Exception as e:
            logger.error(f"Report export error: {str(e)}")
            QMessageBox.critical(self, "Error", f"Failed to export report:\n{str(e)}")
    
    def show_error(self, message):
        """Show error message"""
        QMessageBox.critical(self, "Error", message)
        self.status_label.setText("Analysis Failed")
        self.progress_bar.setValue(0)


if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = MobileMoneyAnalyzer(lambda: None)  # Placeholder callback for standalone testing
    window.show()
    sys.exit(app.exec_())