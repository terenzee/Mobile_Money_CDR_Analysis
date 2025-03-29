import os
import sys
import pandas as pd
import folium
from folium.plugins import HeatMap
import webbrowser
import matplotlib.pyplot as plt
import matplotlib
matplotlib.use('Agg')
from geopy.geocoders import Nominatim
from geopy.exc import GeocoderTimedOut, GeocoderUnavailable
from docx import Document
from docx.shared import Inches, Pt
import datetime
import math
import logging
from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout, QPushButton, 
    QLabel, QFrame, QFileDialog, QMessageBox, QTreeWidget, QTreeWidgetItem, 
    QProgressBar, QScrollArea, QLineEdit, QSplitter, QStackedWidget, QTabWidget
)
from PyQt5.QtCore import Qt, QThread, pyqtSignal
from PyQt5.QtGui import QPixmap
import networkx as nx
from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg as FigureCanvas
from matplotlib.figure import Figure
import seaborn as sns
from pyvis.network import Network
import tempfile
import subprocess

# Suppress libpng warning
os.environ['QT_LOGGING_RULES'] = '*.warning=false'

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('mtn_cdr_analyzer.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# THEME COLORS
BACKGROUND_COLOR = "#1E1E2F"  # Dark blue background
SIDEBAR_COLOR = "#1b2944"     # Sidebar background
CONTENT_BG_COLOR = "#2a3a5a"  # Content area background
TEXT_COLOR = "#FFFFFF"        # White text
ACCENT_COLOR = "#feca18"      # MTN yellow
ACCENT_TEXT = "#0c0d41"       # MTN dark blue for text on yellow
HIGHLIGHT_COLOR = "#e6b800"   # Slightly darker yellow on hover
ERROR_COLOR = "#FF6B6B"       # Light red for errors
SUCCESS_COLOR = "#4CAF50"     # Green for success states
BORDER_COLOR = "#3a4a6a"      # Border color

APP_STYLESHEET = f"""
    /* Main window */
    QMainWindow {{
        background-color: {BACKGROUND_COLOR};
        color: {TEXT_COLOR};
    }}
    
    /* Sidebar */
    QFrame#sidebar {{
        background-color: {SIDEBAR_COLOR};
        border-right: 1px solid {BORDER_COLOR};
    }}
    
    /* Content area */
    QWidget#content_area {{
        background-color: {BACKGROUND_COLOR};
        padding: 20px;
    }}
    
    /* Tree widget */
    QTreeWidget {{
        background-color: {CONTENT_BG_COLOR};
        color: {TEXT_COLOR};
        border: 1px solid {BORDER_COLOR};
        font-size: 12px;
    }}
    
    QTreeWidget::item {{
        padding: 5px;
    }}
    
    QHeaderView::section {{
        background-color: {CONTENT_BG_COLOR};
        color: {TEXT_COLOR};
        padding: 5px;
        border: none;
    }}
    
    /* Standard Buttons */
    QPushButton {{
        background-color: {ACCENT_COLOR};
        color: {ACCENT_TEXT};
        font-size: 14px;
        font-weight: bold;
        padding: 8px 12px;
        border: none;
        border-radius: 5px;
        min-width: 100px;
        min-height: 40px;
    }}
    
    QPushButton:hover {{
        background-color: {HIGHLIGHT_COLOR};
    }}
    
    /* Sidebar Buttons */
    QFrame#sidebar QPushButton {{
        min-width: 180px;
        min-height: 45px;
    }}
    
    /* Important Buttons */
    QPushButton.important {{
        min-width: 150px;
        padding: 10px 15px;
    }}
    
    /* Reset Button */
    QPushButton.reset {{
        background-color: {ERROR_COLOR};
        color: white;
    }}
    
    QPushButton.reset:hover {{
        background-color: #C0392B;
    }}
    
    /* Progress bar */
    QProgressBar {{
        background: {CONTENT_BG_COLOR};
        border: 1px solid {BORDER_COLOR};
        text-align: center;
        color: {TEXT_COLOR};
    }}
    
    QProgressBar::chunk {{
        background-color: {ACCENT_COLOR};
    }}
    
    /* Line edits */
    QLineEdit {{
        background: {CONTENT_BG_COLOR};
        border: 1px solid {BORDER_COLOR};
        padding: 5px;
        color: {TEXT_COLOR};
    }}
    
    /* Labels */
    QLabel {{
        color: {TEXT_COLOR};
    }}
    
    QLabel#file_label {{
        qproperty-alignment: AlignCenter;
        font-size: 14px;
    }}
    
    QLabel#file_label[error=true] {{
        color: {ERROR_COLOR};
    }}
    
    QLabel#file_label[success=true] {{
        color: {SUCCESS_COLOR};
    }}
    
    /* Scroll areas */
    QScrollArea {{
        border: none;
    }}
    
    QScrollBar:vertical {{
        background: {CONTENT_BG_COLOR};
        width: 10px;
    }}
    
    QScrollBar::handle:vertical {{
        background: {BORDER_COLOR};
        min-height: 20px;
    }}
    
    /* Logo frame - matches sidebar background */
    QFrame#logo_frame {{
        background-color: {SIDEBAR_COLOR};
        padding: 5px;
        margin-top: 10px;
    }}
    
    /* Tab widget */
    QTabWidget::pane {{
        border: 1px solid {BORDER_COLOR};
    }}
    
    QTabBar::tab {{
        background: {CONTENT_BG_COLOR};
        color: {TEXT_COLOR};
        padding: 8px;
        border: 1px solid {BORDER_COLOR};
        border-bottom: none;
    }}
    
    QTabBar::tab:selected {{
        background: {ACCENT_COLOR};
        color: {ACCENT_TEXT};
    }}
"""

class AnalysisThread(QThread):
    update_progress = pyqtSignal(str, int)
    analysis_complete = pyqtSignal(list, dict)  # Includes visualization paths and insights
    error_occurred = pyqtSignal(str)

    def __init__(self, df):
        super().__init__()
        self.df = df
        self.visualizations = {}  # To store paths to generated visualizations
        self.insights = []       # To store analysis insights
        self.geolocator = Nominatim(user_agent="mtn_cdr_analyzer")
        self.geolocation_cache = {}

    def run(self):
        try:
            if self.df is None:
                self.error_occurred.emit("No data loaded for analysis")
                return

            results = []
            total_steps = 15  # Increased for new features
            current_step = 0

            # Clean phone numbers and IMEIs to remove decimal points
            self.df['calling_no'] = self.df['calling_no'].astype(str).str.replace(r'\.0$', '', regex=True)
            self.df['called_no'] = self.df['called_no'].astype(str).str.replace(r'\.0$', '', regex=True)
            self.df['imei'] = self.df['imei'].astype(str).str.replace(r'\.0$', '', regex=True)

            # Basic call analysis
            self.update_progress.emit("Analyzing call patterns...", int(current_step/total_steps*100))
            
            calls = self.df
            total_calls = len(calls)
            unique_callers = calls["calling_no"].nunique()
            top_contacts = calls["called_no"].value_counts().head(5)
            total_duration = calls["duration"].sum()
            avg_duration = calls["duration"].mean()

            # Get last call details
            last_call = calls.iloc[-1]
            last_call_details = {
                'time': last_call["event_date_time"],
                'location': f"Lat: {last_call['latitude']:.5f}, Lon: {last_call['longitude']:.5f}",
                'address': self.get_location_name(last_call['latitude'], last_call['longitude']),
                'calling': last_call["calling_no"],
                'called': last_call["called_no"],
                'duration': f"{last_call['duration']:.1f} sec",
                'type': last_call.get("call_type", "VOICE")  # Assuming there's a call_type column
            }

            results.extend([
                ("Total Calls", f"{total_calls}"),
                ("Unique Callers", f"{unique_callers}"),
                ("Total Call Duration (sec)", f"{total_duration}"),
                ("Average Call Duration (sec)", f"{avg_duration:.2f}"),
                ("Top 5 Contacts", "Frequency")
            ])

            for contact, freq in top_contacts.items():
                results.append((f"{contact}", f"{int(freq)}"))  # Ensure frequency is integer

            current_step += 1
            self.update_progress.emit("Analyzing device information...", int(current_step/total_steps*100))

            # Device analysis
            imei_counts = calls["imei"].value_counts().head(5)
            results.append(("Top 5 IMEIs", "Frequency"))
            for imei, freq in imei_counts.items():
                results.append((f"{imei}", f"{int(freq)}"))  # Ensure frequency is integer

            current_step += 1
            self.update_progress.emit("Analyzing geolocations...", int(current_step/total_steps*100))

            # Enhanced geolocation analysis with time periods
            calls['event_date_time'] = pd.to_datetime(calls['event_date_time'])
            calls['hour'] = calls['event_date_time'].dt.hour
            calls['day_of_week'] = calls['event_date_time'].dt.day_name()
            calls['period'] = calls['hour'].apply(self.get_time_period)
            
            # Get most frequent locations for each period
            period_locations = {}
            for period in ['Night', 'Morning', 'Afternoon', 'Evening']:
                period_data = calls[calls['period'] == period]
                if not period_data.empty:
                    period_counts = period_data.groupby(["latitude", "longitude"]).size().reset_index(name="count")
                    freq_location = period_counts.loc[period_counts['count'].idxmax()]
                    period_locations[period] = {
                        'lat': freq_location["latitude"],
                        'lon': freq_location["longitude"],
                        'count': int(freq_location["count"]),
                        'address': self.get_location_name(freq_location["latitude"], freq_location["longitude"])
                    }

            current_step += 1
            self.update_progress.emit("Generating advanced visualizations...", int(current_step/total_steps*100))

            # Generate all visualizations
            self.generate_call_distribution_charts(calls)
            self.visualizations['hourly'] = "hourly_distribution.png"
            self.visualizations['daily'] = "day_distribution.png"
            self.generate_location_plot(calls)
            self.visualizations['locations'] = "location_plot.png"
            self.generate_map_with_features(calls)
            self.visualizations['map'] = "MTN_CDR_Map.html"
            self.generate_network_graph(calls)
            self.visualizations['network'] = "call_network.html"
            self.generate_statistical_charts(calls)
            self.visualizations['stats'] = "statistical_analysis.png"

            current_step += 1
            self.update_progress.emit("Finalizing results...", int(current_step/total_steps*100))

            # Add last call details to results
            results.extend([
                ("\nLast Call Information", ""),
                ("Time", str(last_call_details['time'])),
                ("Location", last_call_details['location']),
                ("Address", last_call_details['address']),
                ("Calling Number", last_call_details['calling']),
                ("Called Number", last_call_details['called']),
                ("Duration", last_call_details['duration']),
                ("Type", last_call_details['type'])
            ])

            # Add period locations to results
            results.append(("\nMost Frequent Locations by Time Period", ""))
            for period, loc in period_locations.items():
                results.extend([
                    (f"{period} Location", f"Lat: {loc['lat']:.5f}, Lon: {loc['lon']:.5f}"),
                    (f"{period} Calls", f"{loc['count']}"),
                    (f"{period} Address", loc['address'])
                ])

            # Generate insights
            self.generate_insights(calls, period_locations, last_call_details)

            self.analysis_complete.emit(results, {
                'visualizations': self.visualizations,
                'insights': self.insights,
                'period_locations': period_locations,
                'last_call': last_call_details
            })

        except Exception as e:
            logger.error(f"Analysis error: {str(e)}")
            self.error_occurred.emit(f"Analysis failed: {str(e)}")

    def get_time_period(self, hour):
        """Categorize time into periods"""
        if 0 <= hour < 6:
            return 'Night'
        elif 6 <= hour < 12:
            return 'Morning'
        elif 12 <= hour < 18:
            return 'Afternoon'
        else:
            return 'Evening'

    def get_location_name(self, latitude, longitude):
        cache_key = (latitude, longitude)
        if cache_key in self.geolocation_cache:
            return self.geolocation_cache[cache_key]

        try:
            location = self.geolocator.reverse((latitude, longitude), exactly_one=True, timeout=10)
            address = location.address if location else "Unknown Location"
            self.geolocation_cache[cache_key] = address
            return address
        except (GeocoderTimedOut, GeocoderUnavailable):
            return "Location Unavailable"

    def generate_call_distribution_charts(self, calls):
        """Generate time distribution charts"""
        # Hourly distribution
        plt.figure(figsize=(10, 5))
        calls['hour'].value_counts().sort_index().plot(kind='bar')
        plt.title("Call Distribution by Hour")
        plt.xlabel("Hour of Day")
        plt.ylabel("Number of Calls")
        plt.savefig("hourly_distribution.png", dpi=100, bbox_inches='tight')
        plt.close()
        
        # Day of week distribution
        plt.figure(figsize=(8, 5))
        calls['day_of_week'].value_counts().plot(kind='bar')
        plt.title("Call Distribution by Day of Week")
        plt.xlabel("Day of Week")
        plt.ylabel("Number of Calls")
        plt.savefig("day_distribution.png", dpi=100, bbox_inches='tight')
        plt.close()

    def generate_location_plot(self, calls):
        """Generate geolocation visualization"""
        plt.figure(figsize=(10, 6))
        calls = calls.dropna(subset=['latitude', 'longitude'])
        plt.scatter(calls["longitude"], calls["latitude"], alpha=0.6, edgecolors="w", s=100)
        plt.xlabel("Longitude")
        plt.ylabel("Latitude")
        plt.title("Call Locations")
        plt.grid(True)
        plt.savefig("location_plot.png", dpi=100, bbox_inches='tight')
        plt.close()

    def generate_map_with_features(self, calls):
        """Generate enhanced map with all features"""
        try:
            calls = calls.dropna(subset=["latitude", "longitude"])
            if calls.empty:
                raise ValueError("No valid location data available")
            
            location_counts = calls.groupby(["latitude", "longitude", "azimuth"]).size().reset_index(name="count")
            highest_location = location_counts.loc[location_counts['count'].idxmax()]
            lat, lon, azimuth = highest_location["latitude"], highest_location["longitude"], highest_location["azimuth"]
            
            m = folium.Map(location=[lat, lon], zoom_start=12)
            
            # Add cell tower markers
            for _, row in location_counts.iterrows():
                folium.Marker(
                    [row["latitude"], row["longitude"]],
                    popup=f"Address: {self.get_location_name(row['latitude'], row['longitude'])}\nCalls: {int(row['count'])}\nAzimuth: {int(row['azimuth'])}°",
                    icon=folium.Icon(color='blue', icon='tower-cell', prefix='fa')
                ).add_to(m)
                
                # Add azimuth direction indicator
                azimuth_rad = math.radians(row["azimuth"])
                end_lat = row["latitude"] + (0.6 / 111.32) * math.cos(azimuth_rad)
                end_lon = row["longitude"] + (0.6 / (111.32 * math.cos(math.radians(row["latitude"])))) * math.sin(azimuth_rad)
                folium.PolyLine(
                    locations=[[row["latitude"], row["longitude"]], [end_lat, end_lon]],
                    color='red',
                    weight=2
                ).add_to(m)
            
            # Highlight most frequent location
            folium.Circle(
                location=[lat, lon],
                radius=400,
                color='red',
                fill=True,
                fill_color='red'
            ).add_to(m)
            
            folium.Marker(
                [lat, lon],
                popup=f"Address: {self.get_location_name(lat, lon)}\nCalls: {int(highest_location['count'])}\nAzimuth: {int(highest_location['azimuth'])}°",
                icon=folium.Icon(color='red', icon='tower-cell', prefix='fa')
            ).add_to(m)
            
            # Add frequent locations for each period
            colors = {"Night": "blue", "Morning": "green", "Afternoon": "orange", "Evening": "purple"}
            for period in colors.keys():
                period_data = calls[calls['period'] == period]
                if not period_data.empty:
                    period_counts = period_data.groupby(["latitude", "longitude"]).size().reset_index(name="count")
                    freq_location = period_counts.loc[period_counts['count'].idxmax()]
                    freq_lat, freq_lon = freq_location["latitude"], freq_location["longitude"]
                    
                    # Get azimuth for this location
                    azimuth = period_data[(period_data['latitude'] == freq_lat) & 
                                         (period_data['longitude'] == freq_lon)]['azimuth'].mode()[0]
                    
                    folium.Marker(
                        [freq_lat, freq_lon],
                        popup=f"Most Frequent {period} Location\nAddress: {self.get_location_name(freq_lat, freq_lon)}\nCalls: {int(freq_location['count'])}\nAzimuth: {int(azimuth)}°",
                        icon=folium.Icon(color=colors[period], icon='tower-cell', prefix='fa')
                    ).add_to(m)
                    
                    # Add azimuth line
                    azimuth_rad = math.radians(azimuth)
                    end_lat = freq_lat + (0.6 / 111.32) * math.cos(azimuth_rad)
                    end_lon = freq_lon + (0.6 / (111.32 * math.cos(math.radians(freq_lat)))) * math.sin(azimuth_rad)
                    folium.PolyLine(
                        locations=[[freq_lat, freq_lon], [end_lat, end_lon]],
                        color=colors[period],
                        weight=2
                    ).add_to(m)
            
            # Add heatmap
            heat_data = [[row['latitude'], row['longitude'], row['count']] for _, row in location_counts.iterrows()]
            HeatMap(heat_data, radius=15).add_to(m)
            
            # Add legend
            legend_html = '''
            <div style="position: fixed; bottom: 50px; left: 50px; width: 200px; height: 180px; 
                        border:2px solid grey; z-index:9999; font-size:14px; background-color:white;
                        padding: 10px;">
                <p><strong>Period Legend</strong></p>
                <p><span style="color:blue;">⬤</span> Night</p>
                <p><span style="color:green;">⬤</span> Morning</p>
                <p><span style="color:orange;">⬤</span> Afternoon</p>
                <p><span style="color:purple;">⬤</span> Evening</p>
                <p style="font-size:8px; margin-top:10px;">
                    MTN CDR Analyzer | Developed by Terence | © 2025
                </p>
            </div>
            '''
            m.get_root().html.add_child(folium.Element(legend_html))
            
            # Save map
            m.save("MTN_CDR_Map.html")
            
        except Exception as e:
            logger.error(f"Map generation error: {str(e)}")
            raise

    def generate_network_graph(self, calls):
        """Generate interactive network graph of call patterns"""
        try:
            # Create a network graph of calls
            top_contacts = calls["called_no"].value_counts().head(20)
            filtered_calls = calls[calls["called_no"].isin(top_contacts.index)]
            
            # Create graph
            G = nx.Graph()
            
            # Add nodes (callers and called numbers)
            for num in filtered_calls["calling_no"].unique():
                G.add_node(num, size=10, title=num, group=1)
                
            for num in filtered_calls["called_no"].unique():
                G.add_node(num, size=5, title=num, group=2)
            
            # Add edges (calls between numbers)
            call_counts = filtered_calls.groupby(["calling_no", "called_no"]).size().reset_index(name="weight")
            for _, row in call_counts.iterrows():
                G.add_edge(row["calling_no"], row["called_no"], weight=row["weight"], title=f"{int(row['weight'])} calls")
            
            # Generate interactive visualization
            net = Network(height="600px", width="100%", bgcolor="#222222", font_color="white")
            net.from_nx(G)
            
            # Save to HTML file
            net.save_graph("call_network.html")
            
        except Exception as e:
            logger.error(f"Network graph generation error: {str(e)}")
            raise

    def generate_statistical_charts(self, calls):
        """Generate statistical analysis charts"""
        try:
            plt.figure(figsize=(12, 8))
            
            # Create subplots
            fig, axes = plt.subplots(2, 2, figsize=(12, 10))
            
            # Duration distribution
            sns.histplot(calls["duration"], bins=30, kde=True, ax=axes[0, 0])
            axes[0, 0].set_title("Call Duration Distribution")
            axes[0, 0].set_xlabel("Duration (seconds)")
            axes[0, 0].set_ylabel("Frequency")
            
            # Calls by period
            period_counts = calls['period'].value_counts()
            period_counts.plot(kind='bar', ax=axes[0, 1], color=ACCENT_COLOR)
            axes[0, 1].set_title("Calls by Time Period")
            axes[0, 1].set_xlabel("Time Period")
            axes[0, 1].set_ylabel("Number of Calls")
            
            # Top called numbers
            top_called = calls['called_no'].value_counts().head(10)
            top_called.plot(kind='barh', ax=axes[1, 0], color=ACCENT_COLOR)
            axes[1, 0].set_title("Top 10 Called Numbers")
            axes[1, 0].set_xlabel("Number of Calls")
            
            # Call duration by period
            sns.boxplot(x='period', y='duration', data=calls, ax=axes[1, 1])
            axes[1, 1].set_title("Call Duration by Time Period")
            axes[1, 1].set_xlabel("Time Period")
            axes[1, 1].set_ylabel("Duration (seconds)")
            
            plt.tight_layout()
            plt.savefig("statistical_analysis.png", dpi=100, bbox_inches='tight')
            plt.close()
            
        except Exception as e:
            logger.error(f"Statistical charts generation error: {str(e)}")
            raise

    def generate_insights(self, calls, period_locations, last_call):
        """Generate enhanced insights"""
        self.insights = []
        
        # Time patterns insight
        peak_hour = calls['hour'].value_counts().idxmax()
        self.insights.append(
            f"Peak calling hour: {peak_hour}:00 with {calls['hour'].value_counts().max()} calls"
        )
        
        # Device usage insight
        top_imei = calls['imei'].value_counts().idxmax()
        self.insights.append(
            f"Most used device: IMEI {top_imei} with {calls['imei'].value_counts().max()} calls"
        )
        
        # Call duration insight
        avg_duration = calls['duration'].mean()
        self.insights.append(
            f"Average call duration: {avg_duration:.2f} seconds"
        )
        
        # Frequent contacts insight
        top_contact = calls['called_no'].value_counts().idxmax()
        self.insights.append(
            f"Most contacted number: {top_contact} with {calls['called_no'].value_counts().max()} calls"
        )
        
        # Last call insight
        self.insights.append(
            f"Last call was to {last_call['called']} at {last_call['time']} from location {last_call['location']}"
        )
        
        # Period location insights
        for period, loc in period_locations.items():
            self.insights.append(
                f"Most frequent {period.lower()} location: {loc['address']} with {loc['count']} calls"
            )


class MTNCDRAnalyzer(QMainWindow):
    def __init__(self, back_to_home_callback):
        super().__init__()
        self.back_to_home_callback = back_to_home_callback
        self.setWindowTitle("MTN Call Detail Records Analyzer")
        self.setGeometry(100, 100, 1200, 800)
        self.setStyleSheet(APP_STYLESHEET)
        
        # Initialize variables
        self.filename = None
        self.df = None
        self.doc_path = None
        self.analysis_thread = None
        self.visualizations = {}
        self.insights = []
        
        # Create main widgets
        self.init_ui()
        self.create_home_page()
    
    def init_ui(self):
        """Initialize main UI components"""
        self.central_widget = QWidget()
        self.setCentralWidget(self.central_widget)
        self.main_layout = QHBoxLayout(self.central_widget)
        self.main_layout.setContentsMargins(0, 0, 0, 0)
        self.main_layout.setSpacing(0)
        
        # Create sidebar with styled buttons
        self.sidebar = QFrame()
        self.sidebar.setObjectName("sidebar")
        self.sidebar.setFixedWidth(240)
        self.sidebar_layout = QVBoxLayout(self.sidebar)
        self.sidebar_layout.setAlignment(Qt.AlignTop)
        self.sidebar_layout.setSpacing(10)
        self.sidebar_layout.setContentsMargins(10, 20, 10, 10)
        
        # Add buttons to sidebar
        self.create_sidebar_buttons()
        self.sidebar_layout.addStretch()
        self.add_sidebar_logo()
        
        self.main_layout.addWidget(self.sidebar)
        
        # Create content area
        self.content_stack = QStackedWidget()
        self.main_layout.addWidget(self.content_stack)
    
    def create_sidebar_buttons(self):
        """Create sidebar navigation buttons"""
        buttons = [
            ("Load File", self.load_file, "standard"),
            ("Analyze Data", self.create_analysis_page, "standard"),
            ("Open Map", self.open_map, "standard"),
            ("Export Report", self.export_report, "standard"),
            ("Back to Home", self.back_to_home, "warning")
        ]
        
        for text, callback, btn_type in buttons:
            btn = QPushButton(text)
            btn.setStyleSheet(f"""
                QPushButton {{
                    background-color: {ACCENT_COLOR};
                    color: {ACCENT_TEXT};
                    font-size: 14px;
                    font-weight: bold;
                    padding: 10px;
                    border: none;
                    border-radius: 5px;
                    min-width: 180px;
                    min-height: 45px;
                }}
                QPushButton:hover {{
                    background-color: {HIGHLIGHT_COLOR};
                }}
                QPushButton:disabled {{
                    background-color: #7a7a7a;
                    color: #cccccc;
                }}
            """)
            btn.clicked.connect(callback)
            self.sidebar_layout.addWidget(btn)
            
            if text == "Export Report":
                self.export_button = btn
                self.export_button.setEnabled(False)
    
    def add_sidebar_logo(self):
        """Add MTN logo to sidebar with matching background"""
        logo_frame = QFrame()
        logo_frame.setObjectName("logo_frame")
        logo_frame.setStyleSheet(f"background-color: {SIDEBAR_COLOR}; border: none;")
        logo_layout = QVBoxLayout(logo_frame)
        logo_layout.setAlignment(Qt.AlignCenter)
        logo_layout.setContentsMargins(5, 5, 5, 5)
        
        logo_label = QLabel()
        logo_path = os.path.join("assets", "mtn_logo.png")
        if os.path.exists(logo_path):
            logo_pixmap = QPixmap(logo_path)
            if not logo_pixmap.isNull():
                # Scale logo to appropriate size
                logo_pixmap = logo_pixmap.scaled(100, 100, Qt.KeepAspectRatio, Qt.SmoothTransformation)
                logo_label.setPixmap(logo_pixmap)
                logo_label.setAlignment(Qt.AlignCenter)
        else:
            # Fallback if logo not found
            logo_label.setText("MTN")
            logo_label.setStyleSheet("""
                font-size: 16px; 
                font-weight: bold; 
                color: white;
            """)
        
        logo_layout.addWidget(logo_label)
        self.sidebar_layout.addWidget(logo_frame)
    
    def back_to_home(self):
        """Reset the app and return to home screen."""
        # Reset all variables
        self.filename = None
        self.df = None
        self.doc_path = None
        self.visualizations = {}
        self.insights = []
        
        # Clear the content
        self.clear_content()
        
        # Reset file label
        self.file_label.setText("📂 No file selected")
        self.file_label.setProperty("error", True)
        self.file_label.style().polish(self.file_label)
        
        # Disable export button
        self.export_button.setEnabled(False)
        
        # Return to home page
        self.content_stack.setCurrentWidget(self.home_page)
        
        # If we have a callback, use it to exit completely
        if self.back_to_home_callback:
            self.back_to_home_callback()
    
    def clear_content(self):
        """Clear the content area"""
        if hasattr(self, 'analysis_page'):
            layout = self.analysis_page.layout()
            if layout:
                while layout.count():
                    item = layout.takeAt(0)
                    widget = item.widget()
                    if widget:
                        widget.deleteLater()
    
    def create_home_page(self):
        """Create the home page"""
        if hasattr(self, 'home_page'):
            self.content_stack.setCurrentWidget(self.home_page)
            return
        
        self.home_page = QWidget()
        home_layout = QVBoxLayout(self.home_page)
        home_layout.setAlignment(Qt.AlignCenter)
        
        title = QLabel("MTN Call Detail Records Analyzer")
        title.setStyleSheet("font-size: 20px; font-weight: bold;")
        title.setAlignment(Qt.AlignCenter)
        home_layout.addWidget(title)
        
        self.file_label = QLabel("📂 No file selected")
        self.file_label.setObjectName("file_label")
        self.file_label.setProperty("error", True)
        self.file_label.style().polish(self.file_label)
        home_layout.addWidget(self.file_label)
        
        home_layout.addStretch()
        self.content_stack.addWidget(self.home_page)
        self.content_stack.setCurrentWidget(self.home_page)
    
    def create_analysis_page(self):
        """Create the analysis page with results"""
        if self.df is None:
            QMessageBox.warning(self, "Error", "Please load a file first!")
            return
        
        if hasattr(self, 'analysis_page'):
            self.content_stack.setCurrentWidget(self.analysis_page)
            return
        
        self.analysis_page = QWidget()
        analysis_layout = QVBoxLayout(self.analysis_page)
        
        title = QLabel("CDR Analysis Report")
        title.setStyleSheet("font-size: 18px; font-weight: bold;")
        title.setAlignment(Qt.AlignCenter)
        analysis_layout.addWidget(title)
        
        search_frame = QFrame()
        search_layout = QHBoxLayout(search_frame)
        
        self.search_input = QLineEdit()
        self.search_input.setPlaceholderText("Search...")
        self.search_input.textChanged.connect(self.filter_treeview)
        search_layout.addWidget(self.search_input)
        
        reset_btn = QPushButton("⭮ Reset")
        reset_btn.setStyleSheet(f"""
            QPushButton {{
                background-color: {ERROR_COLOR};
                color: white;
                font-size: 14px;
                font-weight: bold;
                padding: 8px 12px;
                border: none;
                border-radius: 5px;
            }}
            QPushButton:hover {{
                background-color: #C0392B;
            }}
        """)
        reset_btn.clicked.connect(self.reset_filters)
        search_layout.addWidget(reset_btn)
        
        analysis_layout.addWidget(search_frame)
        
        # Create tab widget for results and visualizations
        self.tab_widget = QTabWidget()
        
        # Results tab
        results_tab = QWidget()
        results_layout = QVBoxLayout(results_tab)
        
        self.tree = QTreeWidget()
        self.tree.setHeaderLabels(["Metric", "Value"])
        self.tree.setColumnWidth(0, 300)
        self.tree.setColumnWidth(1, 500)
        results_layout.addWidget(self.tree)
        
        # Insights section
        insights_label = QLabel("Key Insights")
        insights_label.setStyleSheet("font-size: 16px; font-weight: bold;")
        results_layout.addWidget(insights_label)
        
        self.insights_text = QLabel("Analysis insights will appear here...")
        self.insights_text.setWordWrap(True)
        results_layout.addWidget(self.insights_text)
        
        self.tab_widget.addTab(results_tab, "Results")
        
        # Visualizations tab
        viz_tab = QWidget()
        viz_layout = QVBoxLayout(viz_tab)
        
        self.viz_scroll = QScrollArea()
        self.viz_scroll.setWidgetResizable(True)
        self.viz_container = QWidget()
        self.viz_layout = QVBoxLayout(self.viz_container)
        
        # Placeholder for visualizations
        self.viz_placeholder = QLabel("Visualizations will appear here after analysis")
        self.viz_placeholder.setAlignment(Qt.AlignCenter)
        self.viz_layout.addWidget(self.viz_placeholder)
        
        self.viz_scroll.setWidget(self.viz_container)
        viz_layout.addWidget(self.viz_scroll)
        
        self.tab_widget.addTab(viz_tab, "Visualizations")
        
        analysis_layout.addWidget(self.tab_widget)
        
        self.progress_bar = QProgressBar()
        self.progress_bar.setRange(0, 100)
        analysis_layout.addWidget(self.progress_bar)
        
        self.status_label = QLabel("Analyzing data... Please wait")
        self.status_label.setAlignment(Qt.AlignCenter)
        analysis_layout.addWidget(self.status_label)
        
        self.content_stack.addWidget(self.analysis_page)
        self.content_stack.setCurrentWidget(self.analysis_page)
        
        self.start_analysis()
    
    def start_analysis(self):
        """Start the analysis in a background thread"""
        if self.analysis_thread and self.analysis_thread.isRunning():
            self.analysis_thread.terminate()
        
        self.analysis_thread = AnalysisThread(self.df)
        self.analysis_thread.update_progress.connect(self.update_progress_status)
        self.analysis_thread.analysis_complete.connect(self.display_results)
        self.analysis_thread.error_occurred.connect(self.show_error)
        self.analysis_thread.start()
    
    def update_progress_status(self, message, percent):
        """Update progress status"""
        self.status_label.setText(message)
        self.progress_bar.setValue(percent)
    
    def display_results(self, results, analysis_data):
        """Display analysis results and visualizations"""
        self.tree.clear()
        self.visualizations = analysis_data['visualizations']
        self.insights = analysis_data['insights']
        
        # Display results in tree widget
        for category, details in results:
            item = QTreeWidgetItem([str(category), str(details)])
            self.tree.addTopLevelItem(item)
        
        # Display insights
        self.insights_text.setText("\n• ".join([""] + self.insights))
        
        # Display visualizations
        self.show_visualizations()
        
        # Enable export button
        self.export_button.setEnabled(True)
        
        self.status_label.setText("Analysis Complete")
    
    def show_visualizations(self):
        """Display all generated visualizations with interactive elements"""
        # Clear previous visualizations
        for i in reversed(range(self.viz_layout.count())): 
            self.viz_layout.itemAt(i).widget().setParent(None)
        
        # Remove placeholder if it exists
        if hasattr(self, 'viz_placeholder'):
            self.viz_layout.removeWidget(self.viz_placeholder)
            self.viz_placeholder.deleteLater()
            del self.viz_placeholder
        
        # Create tab widget for visualizations
        viz_tabs = QTabWidget()
        
        # Charts Tab
        charts_tab = QWidget()
        charts_layout = QVBoxLayout(charts_tab)
        charts_scroll = QScrollArea()
        charts_scroll.setWidgetResizable(True)
        charts_container = QWidget()
        charts_inner_layout = QVBoxLayout(charts_container)
        
        # Add each static visualization
        for viz_type in ['hourly', 'daily', 'locations', 'stats']:
            if viz_type in self.visualizations and os.path.exists(self.visualizations[viz_type]):
                # Add title
                title = QLabel(self.get_viz_title(viz_type))
                title.setStyleSheet("font-size: 14px; font-weight: bold;")
                charts_inner_layout.addWidget(title)
                
                # Add image
                pixmap = QPixmap(self.visualizations[viz_type])
                if not pixmap.isNull():
                    label = QLabel()
                    label.setPixmap(pixmap.scaled(800, 600, Qt.KeepAspectRatio, Qt.SmoothTransformation))
                    label.setAlignment(Qt.AlignCenter)
                    charts_inner_layout.addWidget(label)
        
        charts_inner_layout.addStretch()
        charts_scroll.setWidget(charts_container)
        charts_layout.addWidget(charts_scroll)
        viz_tabs.addTab(charts_tab, "Charts")
        
        # Network Graph Tab
        if 'network' in self.visualizations:
            network_tab = QWidget()
            network_layout = QVBoxLayout(network_tab)
            
            # Add description
            desc = QLabel("Interactive Call Network Graph\n(Shows relationships between calling and called numbers)")
            desc.setAlignment(Qt.AlignCenter)
            network_layout.addWidget(desc)
            
            # Add open button
            open_btn = QPushButton("Open Interactive Network Graph")
            open_btn.setStyleSheet(f"""
                QPushButton {{
                    background-color: {ACCENT_COLOR};
                    color: {ACCENT_TEXT};
                    font-size: 14px;
                    font-weight: bold;
                    padding: 10px;
                    min-width: 250px;
                }}
            """)
            open_btn.clicked.connect(lambda: self.open_interactive_viz('network'))
            network_layout.addWidget(open_btn, alignment=Qt.AlignCenter)
            
            # Add static preview
            preview_label = QLabel("Network Graph Preview")
            preview_label.setAlignment(Qt.AlignCenter)
            network_layout.addWidget(preview_label)
            
            # Generate a static preview of the network
            try:
                # Create a simple matplotlib network visualization
                fig = Figure(figsize=(8, 6))
                canvas = FigureCanvas(fig)
                ax = fig.add_subplot(111)
                
                # Create a simple network for preview
                G = nx.Graph()
                sample_nodes = ['Caller', 'Called 1', 'Called 2', 'Called 3']
                G.add_edges_from([('Caller', 'Called 1'), ('Caller', 'Called 2'), ('Caller', 'Called 3')])
                
                nx.draw(G, ax=ax, with_labels=True, node_color=ACCENT_COLOR, 
                       node_size=1000, font_weight='bold')
                ax.set_title("Call Network Concept")
                
                network_layout.addWidget(canvas)
            except Exception as e:
                logger.error(f"Network preview error: {str(e)}")
            
            viz_tabs.addTab(network_tab, "Network Graph")
        
        # Map Tab
        if 'map' in self.visualizations:
            map_tab = QWidget()
            map_layout = QVBoxLayout(map_tab)
            
            # Add description
            desc = QLabel("Interactive Geospatial Map\n(Shows call locations with heatmap and azimuth directions)")
            desc.setAlignment(Qt.AlignCenter)
            map_layout.addWidget(desc)
            
            # Add open button
            open_btn = QPushButton("Open Interactive Map")
            open_btn.setStyleSheet(f"""
                QPushButton {{
                    background-color: {ACCENT_COLOR};
                    color: {ACCENT_TEXT};
                    font-size: 14px;
                    font-weight: bold;
                    padding: 10px;
                    min-width: 250px;
                }}
            """)
            open_btn.clicked.connect(lambda: self.open_interactive_viz('map'))
            map_layout.addWidget(open_btn, alignment=Qt.AlignCenter)
            
            # Add static preview
            preview_label = QLabel("Map Preview")
            preview_label.setAlignment(Qt.AlignCenter)
            map_layout.addWidget(preview_label)
            
            # Add the location plot as preview
            if 'locations' in self.visualizations and os.path.exists(self.visualizations['locations']):
                pixmap = QPixmap(self.visualizations['locations'])
                if not pixmap.isNull():
                    label = QLabel()
                    label.setPixmap(pixmap.scaled(800, 600, Qt.KeepAspectRatio, Qt.SmoothTransformation))
                    label.setAlignment(Qt.AlignCenter)
                    map_layout.addWidget(label)
            
            viz_tabs.addTab(map_tab, "Map")
        
        self.viz_layout.addWidget(viz_tabs)
    
    def open_interactive_viz(self, viz_type):
        """Open interactive visualization in browser"""
        if viz_type not in self.visualizations:
            QMessageBox.warning(self, "Error", f"{viz_type} visualization not available")
            return
        
        viz_path = self.visualizations[viz_type]
        if os.path.exists(viz_path):
            # For HTML files, open in browser
            if viz_path.endswith('.html'):
                webbrowser.open(viz_path)
            else:
                # For other files, try to open with default application
                try:
                    os.startfile(viz_path)  # Windows
                except:
                    try:
                        subprocess.run(["open", viz_path])  # macOS
                    except:
                        subprocess.run(["xdg-open", viz_path])  # Linux
        else:
            QMessageBox.warning(self, "Error", "Visualization file not found")
    
    def filter_treeview(self):
        """Filter tree items based on search text"""
        search_text = self.search_input.text().lower()
        
        for i in range(self.tree.topLevelItemCount()):
            item = self.tree.topLevelItem(i)
            category = item.text(0).lower()
            details = item.text(1).lower()
            item.setHidden(search_text not in category and search_text not in details)
    
    def reset_filters(self):
        """Reset all filters"""
        self.search_input.clear()
        for i in range(self.tree.topLevelItemCount()):
            self.tree.topLevelItem(i).setHidden(False)
    
    def load_file(self):
        """Load Excel file"""
        filename, _ = QFileDialog.getOpenFileName(
            self, "Open File", "", "Excel Files (*.xlsx *.xls);;CSV Files (*.csv)"
        )
        
        if filename:
            try:
                if filename.endswith('.csv'):
                    self.df = pd.read_csv(filename)
                else:
                    self.df = pd.read_excel(filename)

                # Check for required columns
                required_columns = ["calling_no", "called_no", "duration", "event_date_time", "latitude", "longitude", "azimuth"]
                missing_columns = [col for col in required_columns if col not in self.df.columns]
                
                if missing_columns:
                    QMessageBox.critical(self, "Error", f"Missing required columns: {', '.join(missing_columns)}")
                    self.df = None
                    return

                self.filename = filename
                self.file_label.setText(f"📂 Loaded: {os.path.basename(filename)}")
                self.file_label.setProperty("error", False)
                self.file_label.setProperty("success", True)
                self.file_label.style().polish(self.file_label)
                
                # Disable export button until analysis is complete
                self.export_button.setEnabled(False)
                
                QMessageBox.information(self, "Success", "File loaded successfully!")
                
            except Exception as e:
                logger.error(f"File loading error: {str(e)}")
                QMessageBox.critical(self, "Error", f"Failed to load file:\n{str(e)}")
                self.df = None
                self.file_label.setText("📂 Error loading file")
                self.file_label.setProperty("error", True)
                self.file_label.setProperty("success", False)
                self.file_label.style().polish(self.file_label)
    
    def open_map(self):
        """Open generated map in browser"""
        if 'map' not in self.visualizations:
            QMessageBox.warning(self, "Error", "Map not generated yet. Please analyze data first.")
            return
        
        map_path = self.visualizations['map']
        if os.path.exists(map_path):
            webbrowser.open(map_path)
        else:
            QMessageBox.warning(self, "Error", "Map file not found. Please analyze data again.")
    
    def export_report(self):
        """Export comprehensive report to Word document"""
        if not hasattr(self, 'tree') or self.tree.topLevelItemCount() == 0:
            QMessageBox.warning(self, "Error", "No analysis results to export!")
            return
        
        filename, _ = QFileDialog.getSaveFileName(
            self, "Save Report", "", "Word Documents (*.docx)"
        )
        
        if not filename:
            return
        
        try:
            doc = Document()
            
            # Add title and metadata
            doc.add_heading('MTN CDR Analysis Report', 0)
            current_time = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            doc.add_paragraph(f"Report generated on: {current_time}")
            doc.add_paragraph(f"Source file: {os.path.basename(self.filename)}")
            doc.add_paragraph("\n")

            # Add summary statistics
            doc.add_heading('Summary Statistics', level=1)
            for i in range(self.tree.topLevelItemCount()):
                item = self.tree.topLevelItem(i)
                doc.add_paragraph(f"{item.text(0)}: {item.text(1)}")
            doc.add_paragraph("\n")

            # Add insights
            doc.add_heading('Key Insights', level=1)
            for insight in self.insights:
                doc.add_paragraph(insight, style='List Bullet')
            doc.add_paragraph("\n")

            # Add visualizations
            doc.add_heading('Visualizations', level=1)
            
            # Add each visualization that exists in both UI and filesystem
            for viz_type in ['hourly', 'daily', 'locations', 'stats']:
                if viz_type in self.visualizations and os.path.exists(self.visualizations[viz_type]):
                    doc.add_heading(self.get_viz_title(viz_type), level=2)
                    doc.add_picture(self.visualizations[viz_type], width=Inches(6))
                    doc.add_paragraph("\n")

            # Add network graph information
            doc.add_heading('Call Network Analysis', level=1)
            doc.add_paragraph("An interactive network graph has been generated showing the relationships between:")
            doc.add_paragraph("- Calling numbers (larger nodes)", style='List Bullet')
            doc.add_paragraph("- Called numbers (smaller nodes)", style='List Bullet')
            doc.add_paragraph("- Call frequency (edge thickness)", style='List Bullet')
            doc.add_paragraph("\nThe network graph file is saved alongside this report and can be opened in any web browser.")
            
            # Add map information
            doc.add_heading('Geospatial Analysis', level=1)
            doc.add_paragraph("A detailed interactive map has been generated with the following features:")
            doc.add_paragraph("- Cell tower locations with call frequency", style='List Bullet')
            doc.add_paragraph("- Azimuth direction indicators (600m lines)", style='List Bullet')
            doc.add_paragraph("- Most frequent locations for each time period", style='List Bullet')
            doc.add_paragraph("- Heatmap of call density", style='List Bullet')
            doc.add_paragraph("\nThe map file is saved alongside this report and can be opened in any web browser.")
            
            # Add footer
            section = doc.sections[0]
            footer = section.footer
            footer_paragraph = footer.paragraphs[0]
            footer_paragraph.text = "MTN CDR Analysis Report | Developed by Terence | © 2025"
            footer_paragraph.alignment = 1  # Right-align the footer

            # Save document
            doc.save(filename)

            # Try to open the document
            try:
                os.startfile(filename)  # For Windows
            except:
                import subprocess
                subprocess.run(["open", filename])  # For macOS
                subprocess.run(["xdg-open", filename])  # For Linux

            QMessageBox.information(self, "Success", f"Report exported to {filename}")

        except Exception as e:
            logger.error(f"Report export error: {str(e)}")
            QMessageBox.critical(self, "Error", f"Failed to export report:\n{str(e)}")
    
    def show_error(self, message):
        """Show error message"""
        QMessageBox.critical(self, "Error", message)
        self.status_label.setText("Analysis Failed")
        self.progress_bar.setValue(0)
    
    def get_viz_title(self, viz_type):
        """Get display title for visualization type"""
        titles = {
            'hourly': "Hourly Call Distribution",
            'daily': "Daily Call Distribution",
            'locations': "Call Locations",
            'stats': "Statistical Analysis",
            'network': "Call Network Graph",
            'map': "Geospatial Map"
        }
        return titles.get(viz_type, viz_type.replace('_', ' ').title())


if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = MTNCDRAnalyzer(lambda: None)  # Placeholder callback for standalone testing
    window.show()
    sys.exit(app.exec_())